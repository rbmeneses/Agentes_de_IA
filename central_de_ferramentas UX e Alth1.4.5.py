# coding: utf-8
import streamlit as st
import os
import requests
import logging
import base64
import json
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import tempfile
import vertexai
from vertexai.vision_models import ImageGenerationModel
from vertexai.generative_models import GenerativeModel, Part
from PIL import Image as PILImage
from googleapiclient.discovery import build
from google.api_core import exceptions as google_exceptions
# NOVO IMPORT PARA TEXT-TO-SPEECH
from google.cloud import texttospeech
import yaml
from yaml.loader import SafeLoader
import streamlit_authenticator as stauth
from pathlib import Path
from youtube_transcript_api import YouTubeTranscriptApi
from urllib.parse import urlparse, parse_qs



# --- CONFIGURAÇÃO GERAL E DA PÁGINA ---
st.set_page_config(layout="wide", page_title="Minha Plataforma de IA")

# --- 1. DESIGN E ESTILO (UX/UI) ---
# O CSS foi mantido para preservar a identidade visual da aplicação.
st.markdown("""
<style>
    /* Gradiente na barra lateral */
    [data-testid="stSidebar"] {
        background: linear-gradient(135deg, #3B82F6, #8B5CF6);
        color: white;
    }
    [data-testid="stSidebar"] .st-emotion-cache-1629p8f a, [data-testid="stSidebar"] .st-emotion-cache-10trblm {
        color: white;
    }
    /* Estilo dos botões */
    .stButton>button {
        background-color: #8B5CF6;
        color: white;
        border: none;
        border-radius: 8px;
        padding: 10px 20px;
        transition: background-color 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #6D28D9;
    }
    /* Tipografia e Layout */
    body {
        font-family: 'Lato', 'Roboto', sans-serif;
    }
    h1, h2, h3 {
        font-weight: 300; /* Fontes mais leves e elegantes */
    }
    .stTabs, .stRadio {
        margin-top: 2em; /* Espaçamento generoso */
    }
</style>
""", unsafe_allow_html=True)


# --- 2. SISTEMA DE AUTENTICAÇÃO E GESTÃO DE USUÁRIO ---
# A estrutura de autenticação foi mantida, pois é robusta.
config_file = Path(__file__).parent / "config.yaml"
if not config_file.exists():
    # Cria um arquivo de configuração padrão se não existir
    default_config = {
        "credentials": {
            "usernames": {
                "admin": {
                    "email": "admin@example.com",
                    "name": "Administrador",
                    # Senha '12345' - Use uma ferramenta para gerar um hash bcrypt seguro para produção
                    "password": "$2b$12$Yg.i2h.fA94ccbCo3x.iU.W3Yv8M2d2yVpr0dFzgehJe.eAIqfC6C",
                    "api_keys": {
                        "gemini_key": "", "gcp_project_id": "", "gcp_location": "us-central1",
                        "gsearch_key": "", "gsearch_cx": ""
                    }
                }
            }
        },
        "cookie": {"expiry_days": 30, "key": "some_signature_key", "name": "ia_plataforma_cookie"},
        "preauthorized": {"emails": ["admin@example.com"]}
    }
    with open(config_file, 'w') as file:
        yaml.dump(default_config, file, default_flow_style=False)

with open(config_file) as file:
    config = yaml.load(file, Loader=SafeLoader)

authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)

authenticator.login()

name = st.session_state.get("name")
authentication_status = st.session_state.get("authentication_status")
username = st.session_state.get("username")

# --- LÓGICA PRINCIPAL DA APLICAÇÃO ---

if not authentication_status:
    st.warning("Por favor, faça login para acessar a plataforma.")
    if authentication_status is False:
        st.error('Usuário ou senha incorretos.')
    elif authentication_status is None:
        st.info('Bem-vindo! Por favor, insira seu usuário e senha.')

elif authentication_status:
    # --- FUNÇÕES UTILITÁRIAS GLOBAIS E CONSTANTES ---
    GEMINI_API_BASE_URL = "https://generativelanguage.googleapis.com/v1beta"
    GEMINI_FLASH_MODEL = "gemini-1.5-flash"
    GEMINI_PRO_MODEL = "gemini-1.5-pro"
    VERTEX_IMAGE_MODEL = "imagen-3.0-fast-generate-001"

    def get_api_key(key_name):
        """Busca a chave de API do perfil do usuário logado de forma segura."""
        try:
            return config['credentials']['usernames'][username]['api_keys'][key_name]
        except (KeyError, TypeError):
            st.error(f"Chave de API '{key_name}' não encontrada. Configure-a na página 'Perfil e Configurações'.")
            return None

    def save_file_to_user_storage(file_stream, filename):
        """Salva um arquivo na área de armazenamento persistente do usuário."""
        user_dir = Path(__file__).parent / "user_files" / username
        user_dir.mkdir(parents=True, exist_ok=True)
        file_stream.seek(0)
        with open(user_dir / filename, "wb") as f:
            f.write(file_stream.getbuffer())
        st.success(f"Arquivo '{filename}' salvo com sucesso em 'Meus Arquivos'!")

    def get_gemini_response(prompt_text, model_name, temperature, api_key):
        """Envia um prompt para a API Gemini e retorna a resposta em texto."""
        if not api_key:
            st.error("API Key for Gemini is not configured.")
            return None
        api_endpoint = f"{GEMINI_API_BASE_URL}/models/{model_name}:generateContent?key={api_key}"
        payload = {
            "contents": [{"parts": [{"text": prompt_text}]}],
            "generationConfig": {"temperature": temperature},
            "safetySettings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}
            ],
        }
        try:
            response = requests.post(api_endpoint, json=payload, headers={"Content-Type": "application/json"})
            response.raise_for_status()
            result = response.json()
            if 'candidates' in result and result['candidates'][0]['content']['parts'][0]['text']:
                return result['candidates'][0]['content']['parts'][0]['text']
            else:
                st.warning("A resposta da IA não continha o texto esperado.")
                return f"Resposta inesperada da API: {result}"
        except requests.exceptions.RequestException as e:
            st.error(f"Erro de requisição com a API Gemini: {e}")
            return f"Erro ao comunicar com a API Gemini: {e}"
        except Exception as e:
            st.error(f"Erro ao processar o prompt: {e}")
            return f"Erro ao processar o prompt: {e}"

    # --- BARRA LATERAL DE NAVEGAÇÃO ---
    with st.sidebar:
        st.title(f"Bem-vindo, {name}")
        st.markdown("---")

        page = st.radio("Selecione uma Ferramenta:",
                        ("Página Inicial", "Gerador de Exercícios", "Otimizador de Prompt",
                         "Análise Visual de Imagens", "Criador de Aplicativos",
                         "Fábrica de Spritesheets 2D", "Análise de Logs",
                         "Espelho da Mente",
                         "Assistente de Carreira", # RENOMEADO
                         "Narrador IA", # NOVO
                         "Meus Arquivos", "Perfil e Configurações"))

        st.markdown("---")
        authenticator.logout("Logout", "main")

    # --- DEFINIÇÃO DAS PÁGINAS ---
    def page_inicial():
        st.title("🚀 Minhas Ferramentas de IA")
        st.markdown("### Bem-vindo à sua central de ferramentas de Inteligência Artificial.")
        st.success("**Novidade:** Confira o novo **Narrador IA** para transformar texto em áudio e o **Assistente de Carreira** aprimorado com gerador de cartas de apresentação!")
        st.info("Navegue pelas ferramentas usando o menu à esquerda. Configure suas chaves de API na página 'Perfil e Configurações' para habilitar todas as funcionalidades.")

    def page_gerador_exercicios():
        st.header("🧩 Gerador de Exercícios para Estudo Adaptados")
        st.markdown("Crie exercícios a partir de um tema, de um texto completo, ou ajuste questões já existentes.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GEMINI_API_KEY: return

        # As funções auxiliares internas (ge_*) são mantidas como estão
        def ge_extract_image_prompts(content, desired_image_style=""):
            # ... (código original da função mantido) ...
            parts = content.split("[IMAGEM]")
            image_prompts = []
            for i in range(len(parts) - 1):
                before_context = parts[i][-150:]
                after_context = parts[i + 1][:150]
                image_description = f"{before_context.strip()} {after_context.strip()}"
                if desired_image_style:
                    prompt = (f"Uma imagem educacional no estilo '{desired_image_style}' para ilustrar o seguinte conceito: {image_description}. A imagem deve ser clara, didática e focada no objeto principal.")
                else:
                    prompt = f"Uma imagem educacional detalhada para o seguinte contexto: {image_description}"
                image_prompts.append(prompt)
            return image_prompts

        def ge_gerar_imagem_com_vertexai(prompt):
            # ... (código original da função mantido) ...
            if not GCP_PROJECT_ID or not GCP_LOCATION:
                st.error("Project ID e Location do Google Cloud são necessários. Configure no seu perfil.")
                return None
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                modelo_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo_imagem.generate_images(prompt=prompt, number_of_images=1)
                if not resposta or not resposta[0]._image_bytes:
                    st.warning(f"A IA se recusou a gerar a imagem para o prompt: '{prompt}'.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Ocorreu um erro ao gerar a imagem via Vertex AI: {e}")
                return None

        def generate_exercises(data, api_key):
            try:
                gemini_text_model = GEMINI_FLASH_MODEL

                # NOVO: Lógica para diferentes modos de geração
                generation_mode = data['generationMode']
                if generation_mode == "theme":
                    question_prompt = (
                        f"Por favor, gere {data['questionCount']} exercícios do tipo {data['questionType']} para o nível escolar {data['gradeLevel']}, "
                        f"dificuldade {data['difficulty']} sobre o tema '{data['theme']}'.\n"
                    )
                elif generation_mode == "from_text": # NOVO MODO
                     question_prompt = (
                        f"Com base ESTRITAMENTE no texto fornecido abaixo, gere {data['questionCount']} exercícios do tipo {data['questionType']} para o nível escolar {data['gradeLevel']} e dificuldade {data['difficulty']}.\n"
                        f"Para cada exercício, forneça a resposta correta e uma breve justificativa baseada no texto.\n\n"
                        f"--- TEXTO BASE ---\n{data['userText']}\n--- FIM DO TEXTO ---"
                    )

                question_prompt += f"Adapte o conteúdo para {data['specialNeed']}.\n"

                if data['include_images']:
                    question_prompt += ("\n**Instrução Crítica para Imagens:** Você DEVE inserir o marcador `[IMAGEM]` no texto em locais relevantes para ilustrar conceitos-chave. É obrigatório que o marcador `[IMAGEM]` apareça no texto gerado. Exemplo: '...a mitocôndria, que é a usina de energia da célula. [IMAGEM]'\n")

                # O restante da função permanece igual
                # ... (código original da função mantido para gerar o documento) ...
                gemini_payload = {"contents": [{"parts": [{"text": question_prompt}]}]}
                url = f"{GEMINI_API_BASE_URL}/models/{gemini_text_model}:generateContent?key={api_key}"
                headers = {"Content-Type": "application/json"}
                response = requests.post(url, json=gemini_payload, headers=headers)
                response.raise_for_status()
                gemini_result = response.json()
                content = "".join(part.get('text', '') for part in gemini_result['candidates'][0]['content'].get('parts', []))

                images_data = []
                if data['include_images']:
                    image_prompts = ge_extract_image_prompts(content, data['imageStyle'])
                    if image_prompts:
                        with st.spinner(f"Gerando {len(image_prompts)} imagens..."):
                            for prompt_text in image_prompts:
                                img_bytes = ge_gerar_imagem_com_vertexai(prompt_text)
                                if img_bytes: images_data.append(img_bytes)

                file_stream = BytesIO()
                if data['outputFormat'] == 'docx':
                    doc = Document()
                    doc.add_heading('Exercícios Gerados', 0)
                    text_parts = content.split('[IMAGEM]')
                    for i, part_text in enumerate(text_parts):
                        doc.add_paragraph(part_text.strip())
                        if i < len(images_data):
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                                temp_img.write(images_data[i])
                                temp_img_path = temp_img.name
                            doc.add_picture(temp_img_path, width=doc.sections[0].page_width * 0.5)
                            os.remove(temp_img_path)
                    doc.save(file_stream)
                    file_stream.seek(0)
                    return file_stream, 'exercicios.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                else: # Handles txt and pdf
                    if data['outputFormat'] == 'pdf':
                        c = canvas.Canvas(file_stream, pagesize=A4)
                        textobject = c.beginText()
                        textobject.setTextOrigin(10, 750)
                        textobject.setFont("Helvetica", 12)
                        lines = content.split('\n')
                        for line in lines:
                            textobject.textLine(line)
                        c.drawText(textobject)
                        c.showPage()
                        c.save()
                        file_stream.seek(0)
                        return file_stream, f'exercicios.pdf', 'application/pdf'
                    else: # txt
                        file_stream.write(content.encode('utf-8'))
                        file_stream.seek(0)
                        return file_stream, f'exercicios.txt', 'text/plain'

            except Exception as e:
                raise e

        # UI ATUALIZADA com o novo modo
        generation_mode = st.radio(
            "Selecione o modo de geração:",
            ("A partir de um Tema", "A partir de um Texto"),
            key="ge_mode"
        )

        theme_input = ""
        user_text_input = ""

        if generation_mode == "A partir de um Tema":
             theme_input = st.text_area("📝 Tema ou Instruções:", "Fotossíntese para o ensino fundamental", height=100)
        else: # A partir de um Texto
            user_text_input = st.text_area("📄 Cole o texto base aqui:", "Cole aqui um artigo, capítulo ou qualquer material de estudo...", height=200)

        with st.form(key="exercise_form"):
            # Colunas para organização dos controles
            col1, col2, col3 = st.columns(3)
            with col1:
                question_count = st.slider("🔢 Número de Exercícios", 1, 30, 5)
                grade_level = st.selectbox("🎓 Nível Escolar", ["Infantil", "Fundamental I", "Fundamental II", "Médio", "Superior"], index=2)
            with col2:
                difficulty = st.slider("📊 Dificuldade (0-100)", 0, 100, 40)
                question_type = st.selectbox("✍️ Tipo de Questão", ["Discursiva", "Múltipla Escolha"])
            with col3:
                include_images = st.checkbox("🖼️ Incluir imagens com Vertex AI", value=False)
                output_format = st.selectbox("💾 Formato de Saída", ["docx", "pdf", "txt"], help="Escolha DOCX ou PDF para incluir imagens.")
                image_style = st.text_input("🎨 Estilo da Imagem", "desenho vetorial simples", disabled=not include_images)

            special_need = st.selectbox("♿ Necessidade Específica (Opcional)", ["Nenhuma necessidade específica", "Síndrome de Down", "TEA", "Deficiência Intelectual", "TDAH"])

            submit_button = st.form_submit_button("🚀 Gerar Exercícios")

        if submit_button:
            # Validação da entrada
            if generation_mode == "A partir de um Tema" and not theme_input.strip():
                st.error("Por favor, insira um tema.")
            elif generation_mode == "A partir de um Texto" and not user_text_input.strip():
                st.error("Por favor, cole o texto para gerar os exercícios.")
            else:
                with st.spinner("Gerando exercícios..."):
                    try:
                        payload = {
                           'generationMode': 'theme' if generation_mode == "A partir de um Tema" else 'from_text',
                           'theme': theme_input,
                           'userText': user_text_input,
                           'questionCount': question_count,
                           'gradeLevel': grade_level, 'difficulty': difficulty, 'questionType': question_type,
                           'specialNeed': special_need, 'include_images': include_images,
                           'imageStyle': image_style, 'outputFormat': output_format
                        }
                        file_stream, filename, mime = generate_exercises(payload, GEMINI_API_KEY)
                        st.success("✅ Exercícios gerados com sucesso!")
                        c1, c2 = st.columns(2)
                        with c1:
                           st.download_button(label="📥 Baixar Arquivo", data=file_stream, file_name=filename, mime=mime, use_container_width=True)
                        with c2:
                            if st.button("💾 Salvar em Meus Arquivos", use_container_width=True, key="save_exerc"):
                                save_file_to_user_storage(file_stream, filename)
                    except Exception as e:
                        st.error(f"❌ Erro ao gerar os exercícios: {e}")

    # ... (As outras páginas como 'Otimizador de Prompt', 'Análise Visual', etc., são mantidas exatamente como no original) ...
    def page_otimizador_prompt():
        # ... (CÓDIGO ORIGINAL INALTERADO) ...
        st.header("✨ Otimizador e Executor de Prompts com IA")
        st.markdown("Crie um prompt otimizado e use-o imediatamente para gerar texto ou imagens. Você pode incluir uma imagem como contexto.")

        # --- 1. CONFIGURAÇÃO E CHAVES DE API ---
        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')

        # Checa se as chaves essenciais estão presentes
        if not GEMINI_API_KEY:
            st.warning("A chave da API Gemini é necessária. Configure-a em 'Perfil e Configurações'.")
            return

        # --- 2. FUNÇÕES AUXILIARES (INSPIRADAS EM OUTROS MÓDULOS) ---
        def op_analisar_imagem_para_contexto(image_bytes):
            """Usa o modelo multimodal para descrever uma imagem para o contexto do prompt."""
            if not GCP_PROJECT_ID or not GCP_LOCATION:
                st.error("Project ID e Location do Google Cloud são necessários para analisar imagens.")
                return None
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                image_part = Part.from_data(image_bytes, mime_type="image/jpeg")
                prompt_instrucao = "Descreva de forma concisa e objetiva o conteúdo principal desta imagem. Esta descrição será usada como contexto para criar um novo prompt de IA. Foque nos elementos, estilo e composição."
                response = multimodal_model.generate_content([prompt_instrucao, image_part])
                return response.text
            except Exception as e:
                st.error(f"Erro ao analisar a imagem de contexto: {e}")
                return None

        def op_gerar_imagem_com_vertex(prompt):
            """Gera uma imagem usando Vertex AI, similar à fábrica de spritesheets."""
            if not GCP_PROJECT_ID or not GCP_LOCATION:
                st.error("Project ID e Location do Google Cloud são necessários para gerar imagens.")
                return None
            try:
                if 'VERTEX_IMAGE_MODEL' not in globals():
                    st.error("O modelo de imagem Vertex AI não está definido.")
                    return None
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                modelo_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo_imagem.generate_images(prompt=prompt, number_of_images=1, aspect_ratio="1:1")
                if not resposta or not resposta[0]._image_bytes:
                    st.warning("A IA se recusou a gerar a imagem. Tente um prompt diferente.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Erro ao gerar a imagem via Vertex AI: {e}")
                return None

        # --- 3. INTERFACE DO USUÁRIO ---
        media_type = st.selectbox("1. Qual o objetivo final?", ("Gerar Texto", "Gerar Imagem"), key="op_media")

        st.markdown("---")
        st.subheader("2. Forneça o Contexto")

        # Upload de imagem opcional
        uploaded_file = st.file_uploader("Envie uma imagem como inspiração ou contexto (Opcional)", type=["png", "jpg", "jpeg"])
        if uploaded_file:
            st.image(uploaded_file, caption="Imagem de Contexto", width=200)

        request_type = st.selectbox("Tipo de Requisição:", ("Geração de Conteúdo", "Resumo", "Tradução", "Análise de Sentimento", "Ideação", "Geração de Código", "Debug", "Refatoração", "Outro"), key="op_req")
        specific_details = st.text_area("Detalhes Específicos (tom, formato, público-alvo):", placeholder="Ex: Tom formal, formato de lista, para um público de desenvolvedores.", key="op_details")
        content = st.text_area("Requisitos (o que você deseja que a IA faça?):", placeholder="Ex: Escreva um artigo sobre os benefícios da IA para pequenas empresas.", height=150, key="op_content")

        st.markdown("---")
        st.subheader("3. Ajustes da IA")
        gemini_model = st.selectbox("Modelo Gemini para Otimização:", (GEMINI_FLASH_MODEL, GEMINI_PRO_MODEL), key="op_model")
        temperature = st.slider("Criatividade (Temperatura):", 0.0, 1.0, 0.7, 0.05, key="op_temp")

        # --- 4. LÓGICA DE EXECUÇÃO ---
        if st.button("🚀 Otimizar Prompt e Gerar Conteúdo", use_container_width=True):
            if not content.strip():
                st.error("O campo 'Requisitos' é obrigatório.")
            else:
                image_context_description = ""
                # Etapa 0: Analisar imagem se ela existir
                if uploaded_file:
                    with st.spinner("Analisando imagem de contexto..."):
                        image_data = uploaded_file.getvalue()
                        image_context_description = op_analisar_imagem_para_contexto(image_data)
                        if not image_context_description:
                            st.error("Falha ao analisar a imagem. O processo continuará sem o contexto visual.")

                # Etapa 1: Otimizar o prompt
                with st.spinner("Etapa 1: Otimizando seu prompt com IA..."):
                    prompt_for_gemini = (
                        f"Você é um engenheiro de prompts especialista. Sua tarefa é criar um prompt otimizado e detalhado em inglês para uma IA de {'imagem' if media_type == 'Gerar Imagem' else 'texto'}. "
                        f"Use as seguintes informações:\n"
                        f"- Objetivo: {request_type}\n"
                        f"- Detalhes (Tom, formato, público): {specific_details}\n"
                        f"- Requisitos do usuário: {content}\n"
                    )
                    if image_context_description:
                        prompt_for_gemini += f"- Contexto da imagem fornecida: {image_context_description}\n"

                    prompt_for_gemini += "O prompt gerado deve ser direto, claro, rico em detalhes relevantes e pronto para ser usado pela IA de destino."

                    enhanced_prompt = get_gemini_response(prompt_for_gemini, gemini_model, temperature, GEMINI_API_KEY)

                if not enhanced_prompt or "Erro" in enhanced_prompt:
                    st.error(f"Não foi possível gerar o prompt otimizado. Resposta da API: {enhanced_prompt}")
                    return

                with st.expander("Ver Prompt Otimizado Gerado", expanded=True):
                    st.code(enhanced_prompt, language='text')

                # Etapa 2: Usar o prompt otimizado para gerar o conteúdo final
                st.markdown("---")
                st.subheader("✅ Resultado Final")

                with st.spinner(f"Etapa 2: Usando o prompt para gerar a {'imagem' if media_type == 'Gerar Imagem' else 'saída de texto'}..."):
                    if media_type == "Gerar Texto":
                        final_text_result = get_gemini_response(enhanced_prompt, GEMINI_PRO_MODEL, temperature, GEMINI_API_KEY)
                        if final_text_result:
                            st.markdown(final_text_result)
                        else:
                            st.error("Falha ao gerar o texto final.")

                    elif media_type == "Gerar Imagem":
                        if not GCP_PROJECT_ID or not GCP_LOCATION:
                            st.error("Para gerar imagens, o Project ID e a Location do Google Cloud devem ser configurados no seu perfil.")
                            return
                        generated_image_bytes = op_gerar_imagem_com_vertex(enhanced_prompt)
                        if generated_image_bytes:
                            st.image(generated_image_bytes, caption="Imagem gerada com base no prompt otimizado.", use_column_width=True)
                            st.download_button(
                                label="📥 Baixar Imagem (PNG)",
                                data=generated_image_bytes,
                                file_name="imagem_gerada_otimizador.png",
                                mime="image/png",
                                use_container_width=True
                            )
                        else:
                            st.error("Falha ao gerar a imagem final.")
    def page_analise_visual():
        # ... (CÓDIGO ORIGINAL INALTERADO) ...
        st.header("👁️‍🗨️ Análise Visual para Geração de Conteúdo")
        st.markdown("Envie uma imagem (diagrama, UI, esquema, etc.) e a IA irá gerar código e um prompt baseado nela.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GEMINI_API_KEY or not GCP_PROJECT_ID or not GCP_LOCATION: return

        def av_analisar_imagem(image_bytes, prompt_usuario):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL) # Use Pro model for multimodal capabilities
                image_part = Part.from_data(image_bytes, mime_type="image/png")

                instrucao_ia = f"""
                Analise a imagem fornecida. Com base em seu conteúdo visual, gere duas saídas distintas e claramente separadas:

                1.  **Geração de Código:** Crie um trecho de código relevante (Python, HTML, etc.) ou pseudocódigo que represente a lógica, estrutura ou os elementos visuais da imagem. Se for um diagrama, gere a classe correspondente. Se for um layout de UI, gere o HTML/CSS básico. Adicione um comentário no início do bloco de código indicando a linguagem.

                2.  **Geração de Prompt:** Crie um prompt de texto conciso e otimizado que descreva a imagem. Este prompt deve ser formatado para ser usado diretamente na ferramenta 'Otimizador de Prompt', capturando a essência da imagem para uma futura geração de conteúdo.

                Se o usuário fornecer uma instrução específica, leve-a em consideração: '{prompt_usuario}'
                """

                response = multimodal_model.generate_content([instrucao_ia, image_part])
                return response.text

            except google_exceptions.GoogleAPIError as e:
                st.error(f"Google API Error durante análise de imagem: {e.message}")
                return None
            except Exception as e:
                st.error(f"Erro ao analisar a imagem com o modelo multimodal: {e}")
                return None

        uploaded_file = st.file_uploader("Selecione uma imagem", type=["png", "jpg", "jpeg"])
        user_prompt = st.text_input("Instrução Específica (Opcional):", placeholder="Ex: 'Foco no formulário de login', 'Gere o código em Python'")

        if uploaded_file is not None:
            st.image(uploaded_file, caption="Imagem Carregada.", use_column_width=True)

            if st.button("Analisar Imagem e Gerar Conteúdo", use_container_width=True):
                with st.spinner("A IA está analisando a imagem..."):
                    try:
                        image_data = uploaded_file.getvalue()

                        resultado_analise = av_analisar_imagem(image_data, user_prompt)

                        if resultado_analise:
                            st.success("Análise concluída com sucesso!")
                            st.markdown("---")

                            partes = resultado_analise.split("Geração de Prompt:")
                            if len(partes) == 2:
                                st.subheader("Código Gerado a partir da Imagem:")
                                st.code(partes[0].replace("Geração de Código:", "").strip(), language='python') # Default to python, user can change if needed
                                st.subheader("Prompt Otimizado para Descrever a Imagem:")
                                st.code(partes[1].strip(), language='text')
                            else:
                                st.subheader("Saída da IA (formato inesperado):")
                                st.code(resultado_analise, language='text')

                    except Exception as e:
                        st.error(f"Ocorreu um erro no processo: {e}")

    def page_criador_aplicativos():
        st.header("🏗️ Criador de Aplicativos (Scaffolding)")
        st.markdown("Gere a estrutura de arquivos, o código base e até o Dockerfile para suas aplicações.")

        def ca_gerar_scaffold(project_name, app_type, extra_functions=None, include_dockerfile=False):
            templates = {
                "Streamlit App": {
                    "app.py": f"""
import streamlit as st

st.set_page_config(page_title="{project_name}")

def main():
    st.title("Bem-vindo ao {project_name}!")
    st.write("Este é um aplicativo Streamlit gerado pela Plataforma IA Evoluída.")

    name = st.text_input("Qual é o seu nome?")
    if name:
        st.write(f"Olá, {{name}}!")
{extra_functions or ""}
if __name__ == "__main__":
    main()
                    """,
                    "requirements.txt": "streamlit\n"
                },
                "API Flask": {
                    "app.py": f"""
from flask import Flask, jsonify

app = Flask(__name__)

@app.route('/')
def index():
    return jsonify({{"message": "Bem-vindo à API '{project_name}'!"}})

@app.route('/api/data')
def get_data():
    return jsonify({{"id": 1, "name": "Dado de Exemplo"}})
{extra_functions or ""}
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8080)
                    """,
                    "requirements.txt": "Flask\n"
                },
                "Script de Automação": {
                        "main.py": f"""
import os

def main():
    print("Iniciando o script de automação: {project_name}")
    # Exemplo: Listar arquivos no diretório atual
    files = os.listdir('.')
    print("Arquivos encontrados:")
    for f in files:
        print(f"- {{f}}")
    print("Script concluído.")
{extra_functions or ""}
if __name__ == "__main__":
    main()
                        """,
                        "requirements.txt": "", # Scripts podem não ter dependências iniciais
                        "README.md": f"# {project_name}\\n\\nEste é um script de automação gerado pela Plataforma IA."
                }
            }

            scaffold = templates.get(app_type, {})

            # --- LÓGICA PARA ADICIONAR O DOCKERFILE ---
            if include_dockerfile and app_type in ["Streamlit App", "API Flask"]:
                if app_type == "Streamlit App":
                    dockerfile_content = f"""
# Use uma imagem base do Python
FROM python:3.9-slim

# Defina o diretório de trabalho no container
WORKDIR /app

# Copie o arquivo de dependências e instale-as
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

# Copie o resto dos arquivos da aplicação
COPY . .

# Exponha a porta que o Streamlit usa
EXPOSE 8501

# Defina o comando para rodar a aplicação
CMD ["streamlit", "run", "app.py", "--server.port=8501", "--server.address=0.0.0.0"]
"""
                elif app_type == "API Flask":
                    dockerfile_content = f"""
# Use uma imagem base do Python
FROM python:3.9-slim

# Defina o diretório de trabalho no container
WORKDIR /app

# Copie o arquivo de dependências e instale-as
# Adiciona gunicorn para um servidor de produção mais robusto
RUN echo "gunicorn" >> requirements.txt
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

# Copie o resto dos arquivos da aplicação
COPY . .

# Exponha a porta que o Flask/Gunicorn usará
EXPOSE 8080

# Defina o comando para rodar a aplicação com Gunicorn
CMD ["gunicorn", "--bind", "0.0.0.0:8080", "app:app"]
"""
                scaffold['Dockerfile'] = dockerfile_content.strip()

            return scaffold

        st.subheader("1. Defina os Detalhes do Projeto")
        project_name_input = st.text_input("Nome do Projeto:", "MeuNovoApp")
        app_type_select = st.selectbox(
            "Selecione o Tipo de Aplicativo:",
            ("Streamlit App", "API Flask", "Script de Automação")
        )

        st.subheader("2. Adicione Funções Extras com IA")
        extra_func_description = st.text_area(
            "Descreva funções extras para adicionar ao aplicativo (ex: 'uma função que recebe um texto e retorna o resumo'):",
            placeholder="Ex: Adicionar uma calculadora de IMC que recebe peso e altura."
        )

        # --- NOVA OPÇÃO DE UI ---
        st.subheader("3. Opções de Deploy")
        include_dockerfile = st.checkbox("Incluir Dockerfile para deploy?", value=True, help="Gera um arquivo para empacotar sua aplicação com Docker.")


        GEMINI_API_KEY = get_api_key('gemini_key')

        if st.button("Gerar Estrutura do Aplicativo", use_container_width=True):
            if not project_name_input.strip():
                st.error("Por favor, forneça um nome para o projeto.")
            else:
                with st.spinner("Gerando arquivos do projeto..."):
                    extra_functions_code = ""
                    # Gera código extra com IA, se solicitado
                    if extra_func_description.strip():
                        if not GEMINI_API_KEY:
                            st.warning("Chave Gemini API não configurada. Funções extras via IA não serão geradas.")
                        else:
                            st.info("Gerando funções extras com IA...")
                            prompt_for_gemini = (
                                f"Você é um programador Python sênior. Gere o código completo para a seguinte funcionalidade, a ser adicionada em uma aplicação '{app_type_select}'. "
                                f"A função deve ser bem comentada e pronta para uso.\n\n"
                                f"Descrição da Funcionalidade: '{extra_func_description}'"
                            )
                            enhanced_code = get_gemini_response(prompt_for_gemini, GEMINI_PRO_MODEL, 0.5, GEMINI_API_KEY)
                            if enhanced_code:
                                extra_functions_code = "\n\n" + enhanced_code

                    # Gera a estrutura base
                    scaffold_files = ca_gerar_scaffold(project_name_input, app_type_select, extra_functions=extra_functions_code, include_dockerfile=include_dockerfile)

                    if scaffold_files:
                        st.success(f"Estrutura para '{project_name_input}' gerada!")
                        st.balloons()
                        for filename, content in scaffold_files.items():
                            # Determina a linguagem para o realce de sintaxe
                            lang = 'python'
                            if filename.endswith('.txt') or filename.endswith('.md'):
                                lang = 'text'
                            elif filename == 'Dockerfile':
                                lang = 'dockerfile'

                            with st.expander(f"📄 Arquivo: {filename}"):
                                st.code(content, language=lang)
                    else:
                        st.error("Tipo de aplicativo não encontrado.")

    def page_fabrica_spritesheets():
        # ... (CÓDIGO ORIGINAL INALTERADO) ...
        st.header("👹 Fábrica de Spritesheets 2D - Geração por Template Avançado")
        st.markdown("Utiliza um template de prompt de alta qualidade para gerar a spritesheet completa de uma só vez.")

        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GCP_PROJECT_ID or not GCP_LOCATION: return

        def fs_gerar_prompt_template_avancado(data, lista_frames, estilo_arte):
            style_text = ("2D pixel art spritesheet" if estilo_arte == "Pixel Art" else "2D cartoon spritesheet")
            character_text = (f"a powerful '{data['funcao']}' character, wearing a {data['cor']}. "
                              f"Specific details: {data['detalhes']}. Wielding a {data['arma']}.")
            animation_text = ", ".join(lista_frames)
            prompt_final = (f"{style_text}, of {character_text}. "
                            f"The spritesheet must include a complete set of animation frames: {animation_text}. "
                            "Present the sprites in a clear, organized grid layout with a transparent background. "
                            "The character design must be strictly consistent across all poses. "
                            "Use a side-view perspective, optimized for smooth game animation. High resolution image.")
            return prompt_final

        def fs_gerar_imagem_google(prompt):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION) # Initialize Vertex AI here
                st.info(f"Enviando o prompt para a IA...")
                modelo = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo.generate_images(prompt=prompt, number_of_images=1, aspect_ratio="16:9")
                if not resposta or not resposta[0]._image_bytes:
                    st.warning("A IA se recusou a gerar a imagem. Tente uma descrição diferente.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Erro ao gerar a imagem via Vertex AI: {e}")
                return None

        st.subheader("1. Defina os Detalhes do Personagem")
        estilo_arte_selecionado = st.selectbox("Estilo de Arte", ["Pixel Art", "Cartoon"], key="fs_style_select")
        col1, col2 = st.columns(2)
        with col1:
            funcao = st.text_input("Função ou Classe", "Mago Arcano", key="fs_funcao")
            arma = st.text_input("Arma/Equipamento", "cajado de madeira com orbe brilhante", key="fs_arma")
        with col2:
            cor = st.text_input("Cores Dominantes", "manto azul escuro e túnica roxa", key="fs_cor")
        detalhes = st.text_area("Detalhes Visuais", "ancião com longa barba branca e expressão sábia", key="fs_detalhes", height=100)

        st.subheader("2. Selecione as Animações")
        frames_disponiveis = ["Idle", "Walk", "Run", "Jump (take-off and landing)", "Basic Melee Attack (staff swing)", "Magic Spell Cast (fireball)", "Hurt", "Block", "Death animation"]
        frames_selecionados = st.multiselect("Selecione os frames para a spritesheet", options=frames_disponiveis, default=["Idle", "Walk", "Run", "Magic Spell Cast (fireball)"])

        if st.button("👹 Gerar Spritesheet", use_container_width=True):
            if not frames_selecionados:
                st.warning("Selecione pelo menos um frame de animação.")
            else:
                try:
                    with st.spinner("Construindo prompt e gerando a imagem..."):
                        dados_personagem = {"funcao": funcao, "detalhes": detalhes, "arma": arma, "cor": cor}
                        prompt_completo = fs_gerar_prompt_template_avancado(dados_personagem, frames_selecionados, estilo_arte_selecionado)
                        imagem_gerada_bytes = fs_gerar_imagem_google(prompt_completo)

                        if imagem_gerada_bytes:
                            st.success("Spritesheet gerada com sucesso!")
                            st.image(imagem_gerada_bytes, caption="Resultado da geração com o template.", use_container_width=True)
                            st.download_button("📥 Baixar Spritesheet (PNG)", data=imagem_gerada_bytes, file_name=f"{funcao.replace(' ', '_').lower()}_spritesheet.png", mime="image/png", use_container_width=True)
                except Exception as e:
                    st.error(f"Erro ao gerar imagens com Vertex AI: {e}")
    def page_analise_logs():
        # ... (CÓDIGO ORIGINAL INALTERADO) ...
        st.header("📊 Análise de Logs e Eventos com IA")
        st.markdown("Cole o conteúdo do seu log ou faça upload de uma imagem (ex: screenshot de erro) para a IA analisar e fornecer insights.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')

        if not GEMINI_API_KEY or not GCP_PROJECT_ID or not GCP_LOCATION:
            st.warning("Configure suas chaves de API para Gemini, Google Cloud Project ID e Location na aba 'Perfil e Configurações' para usar esta ferramenta.")
            return

        def al_analisar_imagem_multimodal(image_bytes, prompt_usuario):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                image_part = Part.from_data(image_bytes, mime_type="image/png") # Assuming PNG for image logs

                instrucao_ia = f"""
                Analise esta imagem que faz parte de um log ou representa um evento do sistema.
                Extraia qualquer texto visível (OCR), identifique diagramas, gráficos ou estruturas visuais.
                Descreva o conteúdo visual de forma concisa e relevante para a análise de logs.
                Se for um diagrama de fluxo, descreva o fluxo. Se for um erro, transcreva-o.
                Sua análise deve ser útil para complementar a compreensão do log textual.
                Instrução específica do usuário: '{prompt_usuario}'
                """
                response = multimodal_model.generate_content([instrucao_ia, image_part])
                return response.text
            except google_exceptions.GoogleAPIError as e:
                st.error(f"Erro na API do Google ao analisar a imagem: {e.message}")
                return f"Erro na análise visual da imagem: {e.message}"
            except Exception as e:
                st.error(f"Ocorreu um erro inesperado ao analisar a imagem: {e}")
                return f"Erro na análise visual da imagem: {e}"

        with st.form(key="log_analysis_form"):
            st.subheader("1. Forneça o Conteúdo do Log")
            log_content = st.text_area("📜 Cole o conteúdo do log aqui:", height=250,
                                       placeholder="Ex: [2023-10-26 14:30:15] ERROR: Falha na conexão com o banco de dados. Erro: Timeout, host=db.example.com")

            uploaded_image_log = st.file_uploader("🖼️ Ou faça upload de uma imagem de log/evento (opcional):", type=["png", "jpg", "jpeg"])
            image_analysis_instruction = st.text_input("Instrução para análise da imagem (opcional):",
                                                        placeholder="Ex: 'Foque no diagrama de rede', 'Transcreva o erro na tela'")

            st.subheader("2. Detalhes Adicionais (Opcional)")
            col1, col2, col3 = st.columns(3)
            with col1:
                log_type = st.selectbox("Tipo de Log:", ["Geral", "Sistema", "Aplicação", "Segurança", "Rede", "Banco de Dados", "Outro"])
            with col2:
                data_inicio = st.date_input("Data de Início:", value=None, key="log_start_date")
            with col3:
                data_fim = st.date_input("Data de Fim:", value=None, key="log_end_date")

            col4, col5 = st.columns(2)
            with col4:
                gravidade = st.selectbox("Gravidade:", ["Qualquer", "DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"])
            with col5:
                data_source = st.text_input("Fonte de Dados (Sistema/Serviço):", placeholder="Ex: Apache, Nginx, Kubernetes, Servidor DB")

            user_instruction = st.text_area("Instrução específica para a IA (como você quer a análise?):", height=100,
                                            placeholder="Ex: 'Foque em anomalias de desempenho', 'Identifique tentativas de acesso não autorizado'")

            submit_button = st.form_submit_button("🚀 Analisar Log com IA")

        if submit_button:
            if not log_content and not uploaded_image_log:
                st.error("Por favor, cole o conteúdo do log ou faça upload de uma imagem.")
            else:
                with st.spinner("Analisando log e eventos..."):
                    try:
                        image_analysis_result = ""
                        if uploaded_image_log:
                            st.info("Processando imagem de log...")
                            image_bytes = uploaded_image_log.getvalue()
                            image_analysis_result = al_analisar_imagem_multimodal(image_bytes, image_analysis_instruction)
                            if image_analysis_result:
                                st.success("Análise visual concluída.")
                            else:
                                st.warning("Não foi possível obter uma análise visual da imagem.")


                        prompt_parts = []
                        prompt_parts.append("Analise o seguinte log e eventos para identificar padrões, anomalias, problemas potenciais e sugerir melhorias. Forneça um relatório detalhado.")
                        prompt_parts.append(f"\n--- Detalhes do Log ---")
                        prompt_parts.append(f"Tipo de Log: {log_type}")
                        if data_inicio: prompt_parts.append(f"Data de Início: {data_inicio}")
                        if data_fim: prompt_parts.append(f"Data de Fim: {data_fim}")
                        prompt_parts.append(f"Gravidade: {gravidade}")
                        if data_source: prompt_parts.append(f"Fonte de Dados: {data_source}")

                        if log_content:
                            prompt_parts.append(f"\n--- Conteúdo Textual do Log ---\n{log_content}")

                        if image_analysis_result:
                            prompt_parts.append(f"\n--- Análise Visual da Imagem ---\n{image_analysis_result}")

                        if user_instruction:
                            prompt_parts.append(f"\n--- Instruções Específicas do Usuário ---\n{user_instruction}")

                        prompt_parts.append(f"\n--- Estrutura do Relatório Requerida ---")
                        prompt_parts.append("O relatório deve ser estruturado nas seguintes seções: ")
                        prompt_parts.append("1. **Sumário Executivo:** Breve resumo das descobertas mais importantes.")
                        prompt_parts.append("2. **Eventos Chave e Linha do Tempo:** Descrição dos eventos mais críticos ou relevantes, com marcadores de tempo se disponíveis no log.")
                        prompt_parts.append("3. **Padrões Identificados:** Quaisquer repetições, tendências ou comportamentos normais/anormais encontrados.")
                        prompt_parts.append("4. **Anomalias e Problemas:** Detalhes sobre erros, exceções, falhas ou comportamentos inesperados.")
                        prompt_parts.append("5. **Impacto Potencial:** Explicação de como os problemas identificados podem afetar o sistema ou os usuários.")
                        prompt_parts.append("6. **Recomendações e Próximos Passos:** Sugestões acionáveis para resolver problemas, otimizar ou prevenir ocorrências futuras.")
                        prompt_parts.append("Use formatação Markdown para clareza.")

                        full_prompt = "\n".join(prompt_parts)

                        analysis_result = get_gemini_response(full_prompt, GEMINI_PRO_MODEL, temperature=0.2, api_key=GEMINI_API_KEY) # Lower temperature for factual analysis

                        if analysis_result:
                            st.success("✅ Análise de Log Concluída!")
                            st.markdown(analysis_result)
                        else:
                            st.error("❌ Não foi possível gerar a análise do log. Tente novamente ou verifique suas chaves de API.")

                    except Exception as e:
                        st.error(f"❌ Ocorreu um erro inesperado durante a análise: {e}")
    def page_espelho_da_mente():
        # ... (CÓDIGO ORIGINAL INALTERADO) ...
        st.header("✨ Espelho da Mente Dinâmico")
        st.markdown("Transforme pensamentos e sentimentos complexos em arte simbólica.")

        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GCP_PROJECT_ID or not GCP_LOCATION: return

        def em_criar_prompt_avancado_com_ia(dados_usuario):
            GEMINI_API_KEY = get_api_key('gemini_key') # Retrieve API key here as well for this function
            if not GEMINI_API_KEY:
                return None
            modelo_texto = GenerativeModel(GEMINI_FLASH_MODEL)
            prompt_base = f"""
            Sua tarefa é atuar como um engenheiro de prompt especialista em geração de imagens artísticas.
            Converta a descrição do usuário em um prompt em inglês, detalhado e evocativo para um modelo de IA de imagem.
            O objetivo é criar uma metáfora visual do sentimento do usuário.

            **Diretrizes:**
            1. **Estilo e Atmosfera:** O estilo principal é '{dados_usuario['estilo']}'. Descreva a atmosfera usando luz, sombra e a paleta de cores '{dados_usuario['cores']}' para intensificar a emoção.
            2. **Composição:** Cenário: '{dados_usuario['cenario']}'. Ponto focal: '{dados_usuario['objeto']}', que simboliza a experiência.
            3. **Simbolismo:** Incorpore o sentimento '{dados_usuario['pensamento']}' na cena de forma sutil.
            4. **Detalhes Adicionais:** Inclua: '{dados_usuario['detalhes']}'.
            5. **Restrições:** Evite: '{dados_usuario['evitar']}'.

            Gere o prompt final para a IA de imagem.
            """
            try:
                # vertexai.init() is implicitly handled by the library now
                response = modelo_texto.generate_content([prompt_base])
                return response.text.strip()
            except google_exceptions.GoogleAPIError as e:
                st.error(f"Erro na API do Google ao criar prompt para imagem: {e.message}")
                return None
            except Exception as e:
                st.error(f"Ocorreu um erro inesperado ao criar o prompt: {e}")
                return None

        def em_gerar_imagem_google(prompt):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                modelo_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo_imagem.generate_images(prompt=prompt, number_of_images=1, aspect_ratio="1:1")
                if not resposta or not resposta[0]._image_bytes:
                    st.warning("A IA recusou-se a gerar a imagem. Tente uma descrição diferente.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Erro ao gerar a imagem via Vertex AI: {e}")
                return None

        st.subheader("🎨 Construa sua Visão")
        pensamento_usuario = st.text_area("O que você está pensando ou sentindo?", "Nostalgia e esperança por um futuro incerto.", height=100, key="em_thought")
        cenario = st.text_input("Se esse sentimento fosse um lugar, qual seria?", "Uma estação de trem antiga e vazia ao amanhecer.", key="em_place")
        objeto = st.text_input("Qual objeto simboliza melhor este momento?", "Um único broto verde crescendo entre os trilhos.", key="em_object")
        estilo = st.selectbox("Em qual estilo de arte?", ("Surrealismo", "Impressionismo", "Arte Conceitual", "Abstrato", "Aquarela", "Fantasia Sombria", "Cyberpunk"), key="em_style")
        cores = st.text_input("Quais as cores ou paleta?", "Tons de sépia, cinza, com um ponto de luz dourada.", key="em_colors")
        detalhes = st.text_input("Algum detalhe adicional?", "Névoa baixa no chão, raios de sol por uma janela quebrada.", key="em_details_add")
        evitar = st.text_input("O que NÃO incluir na imagem?", "Pessoas, animais, relógios.", key="em_avoid")

        if st.button("✨ Gerar Imagem da Emoção", use_container_width=True):
            if not all([pensamento_usuario, cenario, objeto]):
                st.warning("Por favor, preencha os campos de sentimento, lugar e objeto.")
            else:
                try:
                    with st.spinner("Etapa 1: Interpretando seus sentimentos com IA..."):
                        dados_para_prompt = {"pensamento": pensamento_usuario, "cenario": cenario, "objeto": objeto, "estilo": estilo, "cores": cores, "detalhes": detalhes, "evitar": evitar}
                        prompt_final = em_criar_prompt_avancado_com_ia(dados_para_prompt)

                    if prompt_final:
                        with st.expander("Ver o prompt técnico gerado"):
                            st.write(prompt_final)

                        with st.spinner("Etapa 2: Pintando sua visão..."):
                            imagem_bytes = em_gerar_imagem_google(prompt_final)

                        if imagem_bytes:
                            st.success("Sua imagem foi criada!")
                            st.image(imagem_bytes, caption="Uma reflexão visual do seu sentimento.", use_column_width=True)
                            st.download_button("📥 Baixar Imagem (PNG)", data=imagem_bytes, file_name="espelho_da_mente.png", mime="image/png", use_container_width=True)
                    else:
                        st.error("Não foi possível gerar o prompt para a imagem.")
                except Exception as e:
                    st.error(f"Erro ao gerar a imagem: {e}")

    def page_assistente_carreira():
        st.header("🚀 Assistente de Carreira com IA")
        st.markdown("Encontre vagas, gere cartas de apresentação e analise seu currículo para cada oportunidade.")

        # Recupera as chaves de API necessárias
        GSEARCH_KEY = get_api_key('gsearch_key')
        GSEARCH_CX = get_api_key('gsearch_cx')
        GEMINI_API_KEY = get_api_key('gemini_key')

        if not GSEARCH_KEY or not GSEARCH_CX or not GEMINI_API_KEY:
            st.warning("Configure suas chaves para Google Search e Gemini na aba 'Perfil e Configurações' para usar esta ferramenta.")
            return

        # Função para buscar vagas (inalterada)
        def bv_Google_Search(query, api_key, cx):
            try:
                service = build("customsearch", "v1", developerKey=api_key)
                res = service.cse().list(q=query, cx=cx, num=10).execute()
                return res.get("items", [])
            except Exception as e:
                st.error(f"Erro ao realizar a busca: {e}")
                return []

        # UI para busca de vagas
        query = st.text_input("Digite a profissão ou área que deseja buscar", placeholder="Ex: Engenheiro de Software Python")
        locais = st.text_input("Deseja restringir por local ou site? (Ex: site:linkedin.com)", placeholder="site:linkedin.com OR site:gupy.io")

        if st.button("Buscar Vagas", use_container_width=True):
            if query:
                with st.spinner(f"Buscando vagas para '{query}'..."):
                    full_query_parts = [f'"{query}"']
                    if locais.strip():
                        full_query_parts.append(f"({locais})")
                    else:
                        full_query_parts.append("site:linkedin.com OR site:gupy.io OR site:br.indeed.com")
                    final_dork = " ".join(full_query_parts)
                    st.session_state['job_results'] = bv_Google_Search(final_dork, GSEARCH_KEY, GSEARCH_CX)
                    # Limpa estados de interações anteriores ao buscar novas vagas
                    st.session_state['selected_job_for_resume'] = None

        # Exibe os resultados e os botões de ação
        if 'job_results' in st.session_state and st.session_state['job_results']:
            st.success(f"🔗 {len(st.session_state['job_results'])} resultado(s) encontrados!")

            for index, r in enumerate(st.session_state['job_results']):
                job_key = r.get('link') or f"job_{index}"
                st.markdown(f"### [{r.get('title', 'Vaga sem título')}]({job_key})")
                st.caption(f"🌐 Fonte: {r.get('displayLink', 'N/A')}")
                snippet = r.get('snippet', 'Sem descrição disponível.')
                st.write(snippet)

                col1, col2 = st.columns(2)
                # Botão para gerar carta de apresentação
                with col1:
                    if st.button("📄 Gerar Carta de Apresentação", key=f"cover_letter_{job_key}", use_container_width=True):
                        with st.spinner(f"Criando uma carta de apresentação para a vaga '{r.get('title')}'..."):
                            prompt = (
                                f"Você é um especialista em RH e recrutamento. Crie uma carta de apresentação profissional e convincente em português do Brasil. "
                                f"A carta deve ser baseada nos detalhes da seguinte vaga:\n"
                                f"**Título da Vaga:** {r.get('title')}\n"
                                f"**Descrição/Snippet:** {snippet}\n\n"
                                f"A carta deve ser formal, direta, destacar o interesse na vaga e na empresa, "
                                f"e incluir placeholders como [Seu Nome], [Sua Experiência Chave] e [Nome do Recrutador ou 'Prezados,' para o usuário preencher."
                            )
                            cover_letter = get_gemini_response(prompt, GEMINI_PRO_MODEL, 0.7, GEMINI_API_KEY)
                            st.session_state[f'cover_letter_result_{job_key}'] = cover_letter

                # NOVO: Botão para análise de currículo
                with col2:
                    if st.button("🔎 Analisar Meu Currículo para esta Vaga", key=f"resume_analysis_{job_key}", use_container_width=True):
                        # Define qual vaga foi selecionada para a análise de currículo
                        st.session_state['selected_job_for_resume'] = job_key

                # Exibe a carta de apresentação gerada, se houver
                if f'cover_letter_result_{job_key}' in st.session_state:
                    with st.expander("Ver Carta de Apresentação Gerada", expanded=True):
                        st.markdown(st.session_state[f'cover_letter_result_{job_key}'])

                # NOVO: Exibe a interface de análise de currículo se a vaga foi selecionada
                if st.session_state.get('selected_job_for_resume') == job_key:
                    st.markdown("---")
                    st.subheader(f"Análise de Currículo para: {r.get('title')}")
                    resume_text = st.text_area("Cole o texto completo do seu currículo aqui:", height=250, key=f"resume_text_{job_key}")
                    if st.button("Analisar Agora", key=f"submit_resume_{job_key}", use_container_width=True):
                        if resume_text.strip():
                            with st.spinner("IA analisando seu currículo contra a descrição da vaga..."):
                                resume_prompt = (
                                    f"Você é um coach de carreira e especialista em recrutamento. Analise o currículo fornecido e compare-o com a descrição da vaga abaixo. "
                                    f"Forneça um feedback construtivo e acionável em português do Brasil.\n\n"
                                    f"**DESCRIÇÃO DA VAGA:**\n- Título: {r.get('title')}\n- Detalhes: {snippet}\n\n"
                                    f"**CURRÍCULO DO CANDIDATO:**\n{resume_text}\n\n"
                                    f"**ESTRUTURA DA ANÁLISE (use Markdown):**\n"
                                    f"1.  **Compatibilidade Geral:** Dê uma nota de 0 a 10 e uma breve justificativa.\n"
                                    f"2.  **Pontos Fortes:** Quais partes do currículo se alinham bem com a vaga.\n"
                                    f"3.  **Pontos a Melhorar:** O que está faltando ou pode ser melhorado.\n"
                                    f"4.  **Sugestões de Alteração:** Dê exemplos específicos de como reescrever frases ou seções do currículo para destacar as habilidades pedidas na vaga. "
                                    f"Foque em usar palavras-chave da descrição da vaga."
                                )
                                resume_feedback = get_gemini_response(resume_prompt, GEMINI_PRO_MODEL, 0.5, GEMINI_API_KEY)
                                st.session_state[f'resume_feedback_{job_key}'] = resume_feedback
                        else:
                            st.error("Por favor, cole o texto do seu currículo na caixa acima.")

                # Exibe o feedback da análise de currículo
                if f'resume_feedback_{job_key}' in st.session_state and st.session_state.get('selected_job_for_resume') == job_key:
                     with st.expander("Ver Análise do Currículo", expanded=True):
                        st.markdown(st.session_state[f'resume_feedback_{job_key}'])

                st.markdown("---")

        elif 'job_results' in st.session_state:
             st.warning("Nenhum resultado encontrado. Experimente alterar o termo ou o filtro.")

    def page_narrador_ia():
    # CORREÇÃO: Todo o código abaixo foi indentado
        st.header("🎙️ Narrador IA - Texto para Áudio")
        st.markdown("Transforme texto em narrações ou extraia, traduza e narre o conteúdo de vídeos do YouTube.")

        # Verifica se as credenciais do Google Cloud estão configuradas
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GEMINI_API_KEY = get_api_key('gemini_key')
        if not GCP_PROJECT_ID or not GEMINI_API_KEY:
            st.warning("O ID do Projeto Google Cloud e a Chave Gemini são necessários. Configure no seu perfil.")
            return

    # --- FUNÇÕES AUXILIARES DA PÁGINA ---
        def get_and_translate_transcript(video_url):
            from youtube_transcript_api import YouTubeTranscriptApi
            try:
                video_id = None
                if "watch?v=" in video_url:
                    video_id = video_url.split("watch?v=")[1].split("&")[0]
                elif "youtube.com/watch?v=" in video_url:
                    video_id = video_url.split("watch?v=")[1].split("&")[0]
                elif "youtu.be/" in video_url:
                    video_id = video_url.split("youtu.be/")[1].split("?")[0]

                if not video_id:
                    st.error("URL do YouTube inválida ou formato não reconhecido.")
                    return None

                st.info("Buscando a transcrição do vídeo...")
                transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
                original_transcript = " ".join([item['text'] for item in transcript_list])

                st.info("Tradução em andamento com a IA...")
                translation_prompt = (
                    "Traduza o seguinte texto para o Português do Brasil. Mantenha a estrutura e o significado o mais fiel possível ao original. "
                    "Não adicione nenhuma introdução ou comentário seu, apenas retorne o texto traduzido.\n\n"
                    f"TEXTO ORIGINAL:\n---\n{original_transcript}"
                )

                translated_text = get_gemini_response(
                    translation_prompt,
                    GEMINI_PRO_MODEL,
                    0.2,
                    GEMINI_API_KEY = get_api_key('gemini_key')

                )
                return translated_text

            except Exception as e:
                st.error(f"Não foi possível obter ou traduzir a transcrição: {e}")
                st.warning("Verifique se o vídeo possui legendas (manuais ou automáticas) disponíveis.")
                return None

    # --- INTERFACE COM ABAS ---
        tab1, tab2 = st.tabs(["Digitar Texto", "Narrar Vídeo do YouTube"])

        with tab1:
            st.subheader("Opção 1: Digite ou cole seu texto")
            voices = {
                "Português (Brasil)": {"Feminina": "pt-BR-Standard-A", "Masculina": "pt-BR-Wavenet-B"},
                "Inglês (EUA)": {"Feminina": "en-US-Standard-C", "Masculina": "en-US-Wavenet-D"},
                "Espanhol (Espanha)": {"Feminina": "es-ES-Standard-A", "Masculina": "es-ES-Wavenet-B"}
            }
            text_input_manual = st.text_area("Texto para narrar:", height=200, placeholder="Digite ou cole seu texto aqui...", key="text_manual")

            col1, col2 = st.columns(2)
            with col1:
                language = st.selectbox("Idioma:", list(voices.keys()), key="lang_manual")
            with col2:
                gender = st.selectbox("Gênero da Voz:", list(voices[language].keys()), key="gender_manual")

            voice_name_selected = voices[language][gender]
            speaking_rate = st.slider("Velocidade da Fala:", min_value=0.5, max_value=2.0, value=1.0, step=0.25, key="rate_manual")

            if st.button("🎧 Gerar Áudio do Texto", use_container_width=True, key="btn_manual"):
                if not text_input_manual.strip():
                    st.error("Por favor, insira um texto para gerar o áudio.")
                else:
                    with st.spinner("A IA está gerando sua narração..."):
                        lang_code = voice_name_selected.split('-')[0] + '-' + voice_name_selected.split('-')[1]
                        audio_bytes = ni_sintetizar_fala(text_input_manual, lang_code, voice_name_selected, speaking_rate)
                    if audio_bytes:
                        st.success("Áudio gerado com sucesso!")
                        st.audio(audio_bytes, format="audio/mp3")
                        st.download_button(label="📥 Baixar Áudio (MP3)", data=audio_bytes, file_name="narracao_ia.mp3", mime="audio/mp3")

        with tab2:
            st.subheader("Opção 2: Insira uma URL do YouTube")
            youtube_url = st.text_input("URL do vídeo:", placeholder="https://www.youtube.com/watch?v=...")

            if st.button("Buscar, Traduzir e Narrar", use_container_width=True):
                if youtube_url.strip():
                    with st.spinner("Processo completo iniciado... Isso pode levar alguns instantes."):
                    # Etapa 1: Obter e traduzir o texto
                        translated_text = get_and_translate_transcript(youtube_url)

                        if translated_text:
                            st.success("Transcrição traduzida com sucesso!")
                            st.session_state['translated_transcript_narrador'] = translated_text

                        # Etapa 2: Gerar o áudio a partir do texto traduzido
                            st.info("Gerando a narração do vídeo...")
                        # Vozes fixas para a narração do vídeo para simplificar
                            narration_voice = "pt-BR-Wavenet-B" # Voz masculina padrão
                            narration_rate = 1.0

                            audio_bytes_youtube = ni_sintetizar_fala(translated_text, "pt-BR", narration_voice, narration_rate)

                            if audio_bytes_youtube:
                                st.success("Narração do vídeo gerada!")
                                st.session_state['audio_bytes_youtube'] = audio_bytes_youtube
                            else:
                                st.error("Falha ao gerar o áudio da narração.")
                else:
                    st.warning("Por favor, insira uma URL do YouTube.")

        # Exibe o texto e o áudio gerados se existirem no estado da sessão
            if 'translated_transcript_narrador' in st.session_state:
                with st.expander("Ver Texto Traduzido", expanded=False):
                    st.text_area(
                        "Conteúdo do vídeo em Português:",
                        value=st.session_state.translated_transcript_narrador,
                        height=300,
                        key="edited_transcript_area"
                    )

            if 'audio_bytes_youtube' in st.session_state:
                st.subheader("Áudio da Narração")
                st.audio(st.session_state.audio_bytes_youtube, format="audio/mp3")
                st.download_button(
                    label="📥 Baixar Narração do Vídeo (MP3)",
                    data=st.session_state.audio_bytes_youtube,
                    file_name="narracao_video.mp3",
                    mime="audio/mp3"
             )

    def page_meus_arquivos():
        # ... (CÓDIGO ORIGINAL INALTERADO) ...
        st.header("🗂️ Meus Arquivos")
        st.markdown("Aqui estão os arquivos que você gerou e salvou na plataforma.")

        user_dir = Path(__file__).parent / "user_files" / username
        if not user_dir.exists() or not any(user_dir.iterdir()):
            st.info("Você ainda não salvou nenhum arquivo.")
            return

        files = os.listdir(user_dir)
        if not files: # Check again in case iterdir() returned something but listdir is empty (e.g., hidden files)
                st.info("Você ainda não salvou nenhum arquivo.")
                return

        for file_name in files:
            file_path = user_dir / file_name
            with open(file_path, "rb") as f:
                st.download_button(
                    label=f"📥 Baixar {file_name}",
                    data=f,
                    file_name=file_name,
                    key=f"download_{file_name}"
                )
            st.markdown("---")
    def page_perfil_configuracoes():
        # ... (CÓDIGO ORIGINAL INALTERADO) ...
        st.header("👤 Perfil e Configurações")
        st.subheader("Gerencie suas chaves de API")
        st.info("Suas chaves são armazenadas de forma segura e usadas para alimentar as ferramentas da plataforma.")

        with st.form("api_keys_form"):
            st.write("**Chaves do Google / Gemini**")
            gemini_key = st.text_input("Gemini API Key", value=get_api_key('gemini_key') or "", type="password")
            gcp_project_id = st.text_input("Google Cloud Project ID", value=get_api_key('gcp_project_id') or "")
            gcp_location = st.text_input("Google Cloud Location", value=get_api_key('gcp_location') or "")

            st.write("**Chaves do Google Custom Search**")
            gsearch_key = st.text_input("Google Custom Search API Key", value=get_api_key('gsearch_key') or "", type="password")
            gsearch_cx = st.text_input("Google Custom Search CX ID", value=get_api_key('gsearch_cx') or "")

            submitted = st.form_submit_button("Salvar Configurações")
            if submitted:
                user_credentials = config['credentials']['usernames'][username]
                user_credentials['api_keys']['gemini_key'] = gemini_key
                user_credentials['api_keys']['gcp_project_id'] = gcp_project_id
                user_credentials['api_keys']['gcp_location'] = gcp_location
                user_credentials['api_keys']['gsearch_key'] = gsearch_key
                user_credentials['api_keys']['gsearch_cx'] = gsearch_cx

                with open(config_file, 'w') as file:
                    yaml.dump(config, file, default_flow_style=False)
                st.success("Suas configurações foram salvas com sucesso!")
                st.rerun()

    # --- ROTEADOR DE PÁGINAS ---
    if page == "Página Inicial":
        page_inicial()
    elif page == "Gerador de Exercícios":
        page_gerador_exercicios()
    elif page == "Otimizador de Prompt":
        page_otimizador_prompt()
    elif page == "Análise Visual de Imagens":
        page_analise_visual()
    elif page == "Criador de Aplicativos":
        page_criador_aplicativos()
    elif page == "Fábrica de Spritesheets 2D":
        page_fabrica_spritesheets()
    elif page == "Análise de Logs":
        page_analise_logs()
    elif page == "Espelho da Mente":
        page_espelho_da_mente()
    elif page == "Assistente de Carreira":
        page_assistente_carreira()
    elif page == "Narrador IA":
        page_narrador_ia()
    elif page == "Meus Arquivos":
        page_meus_arquivos()
    elif page == "Perfil e Configurações":
        page_perfil_configuracoes()