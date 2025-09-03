# coding: utf-8
import streamlit as st
import os
import requests
import logging
import base64
import json
import re
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
# --- IMPORTS ADICIONADOS PARA A NOVA FERRAMENTA ---
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.utils import ImageReader
import textwrap
# --- FIM DOS IMPORTS ADICIONADOS ---
import tempfile
import vertexai
from vertexai.vision_models import ImageGenerationModel, Image
from vertexai.generative_models import GenerativeModel, Part
from PIL import Image as PILImage
from googleapiclient.discovery import build
from google.api_core import exceptions as google_exceptions
from google.cloud import texttospeech
import yaml
from yaml.loader import SafeLoader
import streamlit_authenticator as stauth
from pathlib import Path
from urllib.parse import urlparse, parse_qs
# CORREÇÃO DEFINITIVA DOS IMPORTS DO YOUTUBE
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound, VideoUnavailable
# CORREÇÃO DEFINITIVA DOS IMPORTS DO YOUTUBE
#from youtube_transcript_api import YouTubeTranscriptApi
#ytt_api  =  YouTubeTranscriptApi ()


# --- CONFIGURAÇÃO GERAL E DA PÁGINA ---
st.set_page_config(layout="wide", page_title="Minha Plataforma de IA")

# --- 1. DESIGN E ESTILO (UX/UI) ---
st.markdown("""
<style>
    /* Estilos mantidos conforme o original */
    [data-testid="stSidebar"] {
        background: linear-gradient(135deg, #3B82F6, #8B5CF6);
        color: white;
    }
    [data-testid="stSidebar"] .st-emotion-cache-1629p8f a, [data-testid="stSidebar"] .st-emotion-cache-10trblm {
        color: white;
    }
    .stButton>button {
        background-color: #8B5CF6;
        color: white;
        border: none;
        border-radius: 8px;
        padding: 10px 20px;
        transition: background-color 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #6D28D9;
    }
    body {
        font-family: 'Lato', 'Roboto', sans-serif;
    }
    h1, h2, h3 {
        font-weight: 300;
    }
    .stTabs, .stRadio {
        margin-top: 2em;
    }
</style>
""", unsafe_allow_html=True)

# --- 2. SISTEMA DE AUTENTICAÇÃO E GESTÃO DE USUÁRIO ---
config_file = Path(__file__).parent / "config.yaml"
if not config_file.exists():
    default_config = {
        "credentials": {
            "usernames": {
                "admin": {
                    "email": "admin@example.com",
                    "name": "Administrador",
                    "password": "$2b$12$Yg.i2h.fA94ccbCo3x.iU.W3Yv8M2d2yVpr0dFzgehJe.eAIqfC6C",
                    "api_keys": {
                        "gemini_key": "", "gcp_project_id": "", "gcp_location": "us-central1",
                        "gsearch_key": "", "gsearch_cx": ""
                    }
                }
            }
        },
        "cookie": {"expiry_days": 30, "key": "some_signature_key", "name": "ia_plataforma_cookie"},
        "preauthorized": {"emails": ["admin@example.com"]}
    }
    with open(config_file, 'w') as file:
        yaml.dump(default_config, file, default_flow_style=False)

with open(config_file) as file:
    config = yaml.load(file, Loader=SafeLoader)

authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)

authenticator.login()

name = st.session_state.get("name")
authentication_status = st.session_state.get("authentication_status")
username = st.session_state.get("username")

# --- LÓGICA PRINCIPAL DA APLICAÇÃO ---

if not authentication_status:
    st.warning("Por favor, faça login para acessar a plataforma.")
    if authentication_status is False:
        st.error('Usuário ou senha incorretos.')
    elif authentication_status is None:
        st.info('Bem-vindo! Por favor, insira seu usuário e senha.')

elif authentication_status:
    # --- FUNÇÕES UTILITÁRIAS GLOBAIS E CONSTANTES ---
    GEMINI_API_BASE_URL = "https://generativelanguage.googleapis.com/v1beta"
    GEMINI_FLASH_MODEL = "gemini-2.5-flash" 
    GEMINI_PRO_MODEL = "gemini-2.5-pro"
    VERTEX_IMAGE_MODEL = "imagen-3.0-fast-generate-001"
    
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

    def get_api_key(key_name):
        try:
            key = config['credentials']['usernames'][username]['api_keys'][key_name]
            if not key:
                st.error(f"Chave de API '{key_name}' está vazia. Configure-a na página 'Perfil e Configurações'.")
                return None
            return key
        except (KeyError, TypeError):
            st.error(f"Chave de API '{key_name}' não encontrada. Configure-a na página 'Perfil e Configurações'.")
            return None

    def save_file_to_user_storage(file_stream, filename):
        user_dir = Path(__file__).parent / "user_files" / username
        user_dir.mkdir(parents=True, exist_ok=True)
        file_stream.seek(0)
        with open(user_dir / filename, "wb") as f:
            f.write(file_stream.getbuffer())
        st.success(f"Arquivo '{filename}' salvo com sucesso em 'Meus Arquivos'!")

    def get_gemini_response(prompt_text, model_name, temperature, api_key):
        if not api_key:
            logging.error("API Key for Gemini is not configured.")
            return "Erro: A chave da API Gemini não está configurada."
        api_endpoint = f"{GEMINI_API_BASE_URL}/models/{model_name}:generateContent?key={api_key}"
        payload = {
            "contents": [{"parts": [{"text": prompt_text}]}],
            "generationConfig": {"temperature": temperature},
            "safetySettings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}
            ],
        }
        try:
            response = requests.post(api_endpoint, json=payload, headers={"Content-Type": "application/json"})
            response.raise_for_status()
            result = response.json()
            if 'candidates' not in result or not result['candidates']:
                logging.warning(f"Resposta da IA bloqueada ou vazia. Resposta completa: {result}")
                return f"A resposta da IA foi bloqueada por políticas de segurança ou retornou vazia. Detalhes: {result.get('promptFeedback', 'N/A')}"
            return result['candidates'][0]['content']['parts'][0]['text']
        except requests.exceptions.HTTPError as e:
            error_details = e.response.json()
            logging.error(f"Erro HTTP com a API Gemini: {e}. Detalhes: {error_details}")
            return f"Erro de API: {error_details.get('error', {}).get('message', str(e))}"
        except Exception as e:
            logging.error(f"Erro inesperado ao processar o prompt: {e}")
            return f"Erro inesperado ao processar o prompt: {e}"

    # --- BARRA LATERAL DE NAVEGAÇÃO ---
    with st.sidebar:
        st.title(f"Bem-vindo, {name}")
        st.markdown("---")
        page = st.radio("Selecione uma Ferramenta:",
                        ("Página Inicial", "Gerador de Exercícios", "Otimizador de Prompt",
                         "Análise Visual de Imagens", "Criador de Aplicativos",
                         "Fábrica de Spritesheets 2D", "Análise de Logs",
                         "Gerador POP Retrogaming", # <-- PÁGINA ADICIONADA AO MENU
                         "Espelho da Mente",
                         "Pesquisa Avançada (Dorks)",
                         "Narrador IA",
                         "Estúdio de Mistura Visual",
                         "Meus Arquivos", "Perfil e Configurações"))
        st.markdown("---")
        authenticator.logout("Logout", "main")

    # --- DEFINIÇÃO DAS PÁGINAS ---
    def page_inicial():
        st.title("🚀 Minhas Ferramentas de IA")
        st.markdown("### Bem-vindo à sua central de ferramentas de Inteligência Artificial.")
        st.success("**Novidade:** Confira o novo **Narrador IA** para transformar texto em áudio e o **Assistente de Carreira** aprimorado com gerador de cartas de apresentação!")
        st.info("Navegue pelas ferramentas usando o menu à esquerda. Configure suas chaves de API na página 'Perfil e Configurações' para habilitar todas as funcionalidades.")

    def page_gerador_exercicios():
        st.header("🧩 Gerador de Exercícios para Estudo Adaptados")
        st.markdown("Crie exercícios a partir de um tema, de um texto completo, ou ajuste questões já existentes.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GEMINI_API_KEY: return

        # --- FUNÇÕES INTERNAS ---
        def ge_extract_image_prompts(content, desired_image_style=""):
            parts = content.split("[IMAGEM]")
            image_prompts = []
            for i in range(len(parts) - 1):
                before_context = parts[i][-150:]
                after_context = parts[i + 1][:150]
                image_description = f"{before_context.strip()} {after_context.strip()}"
                if desired_image_style:
                    prompt = (f"Uma imagem educacional no estilo '{desired_image_style}' para ilustrar o seguinte conceito: {image_description}. A imagem deve ser clara, didática e focada no objeto principal.")
                else:
                    prompt = f"Uma imagem educacional detalhada para o seguinte contexto: {image_description}"
                image_prompts.append(prompt)
            return image_prompts

        def ge_gerar_imagem_com_vertexai(prompt):
            if not GCP_PROJECT_ID or not GCP_LOCATION:
                st.error("Project ID e Location do Google Cloud são necessários. Configure no seu perfil.")
                return None
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                modelo_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo_imagem.generate_images(prompt=prompt, number_of_images=1)
                if not resposta or not resposta[0]._image_bytes:
                    st.warning(f"A IA se recusou a gerar a imagem para o prompt: '{prompt}'.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Ocorreu um erro ao gerar a imagem via Vertex AI: {e}")
                return None

        def generate_exercises(data, api_key):
            try:
                gemini_text_model = GEMINI_FLASH_MODEL
                generation_mode = data['generationMode']
                if generation_mode == "theme":
                    question_prompt = (f"Por favor, gere {data['questionCount']} exercícios do tipo {data['questionType']} para o nível escolar {data['gradeLevel']}, "
                                       f"dificuldade {data['difficulty']} sobre o tema '{data['theme']}'.\n")
                elif generation_mode == "from_text":
                    question_prompt = (f"Com base ESTRITAMENTE no texto fornecido abaixo, gere {data['questionCount']} exercícios do tipo {data['questionType']} para o nível escolar {data['gradeLevel']} e dificuldade {data['difficulty']}.\n"
                                       f"Para cada exercício, forneça a resposta correta e uma breve justificativa baseada no texto.\n\n"
                                       f"--- TEXTO BASE ---\n{data['userText']}\n--- FIM DO TEXTO ---")
                question_prompt += f"Adapte o conteúdo para {data['specialNeed']}.\n"
                if data['include_images']:
                    question_prompt += ("\n**Instrução Crítica para Imagens:** Você DEVE inserir o marcador `[IMAGEM]` no texto em locais relevantes para ilustrar conceitos-chave. É obrigatório que o marcador `[IMAGEM]` apareça no texto gerado. Exemplo: '...a mitocôndria, que é a usina de energia da célula. [IMAGEM]'\n")
                
                gemini_payload = {"contents": [{"parts": [{"text": question_prompt}]}]}
                url = f"{GEMINI_API_BASE_URL}/models/{gemini_text_model}:generateContent?key={api_key}"
                headers = {"Content-Type": "application/json"}
                response = requests.post(url, json=gemini_payload, headers=headers)
                response.raise_for_status()
                gemini_result = response.json()
                content = "".join(part.get('text', '') for part in gemini_result['candidates'][0]['content'].get('parts', []))
                
                images_data = []
                if data['include_images']:
                    image_prompts = ge_extract_image_prompts(content, data['imageStyle'])
                    if image_prompts:
                        with st.spinner(f"Gerando {len(image_prompts)} imagens..."):
                            for prompt_text in image_prompts:
                                img_bytes = ge_gerar_imagem_com_vertexai(prompt_text)
                                if img_bytes: images_data.append(img_bytes)
                
                file_stream = BytesIO()
                if data['outputFormat'] == 'docx':
                    doc = Document()
                    doc.add_heading('Exercícios Gerados', 0)
                    text_parts = content.split('[IMAGEM]')
                    for i, part_text in enumerate(text_parts):
                        doc.add_paragraph(part_text.strip())
                        if i < len(images_data):
                            image_stream = BytesIO(images_data[i])
                            doc.add_picture(image_stream, width=doc.sections[0].page_width * 0.5)
                    doc.save(file_stream)
                    file_stream.seek(0)
                    return file_stream, 'exercicios.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif data['outputFormat'] == 'pdf':
                    c = canvas.Canvas(file_stream, pagesize=A4)
                    textobject = c.beginText()
                    textobject.setTextOrigin(10, 750)
                    textobject.setFont("Helvetica", 12)
                    lines = content.split('\n')
                    for line in lines:
                        textobject.textLine(line)
                    c.drawText(textobject)
                    c.showPage()
                    c.save()
                    file_stream.seek(0)
                    return file_stream, 'exercicios.pdf', 'application/pdf'
                else: # txt
                    file_stream.write(content.encode('utf-8'))
                    file_stream.seek(0)
                    return file_stream, 'exercicios.txt', 'text/plain'
            except Exception as e:
                raise e

        # --- INTERFACE DO USUÁRIO (UI) ---
        generation_mode = st.radio("Selecione o modo de geração:", ("A partir de um Tema", "A partir de um Texto"), key="ge_mode")
        theme_input = ""
        user_text_input = ""
        if generation_mode == "A partir de um Tema":
            theme_input = st.text_area("📝 Tema ou Instruções:", "Fotossíntese para o ensino fundamental", height=100)
        else:
            user_text_input = st.text_area("📄 Cole o texto base aqui:", "Cole aqui um artigo, capítulo ou qualquer material de estudo...", height=200)
        
        with st.form(key="exercise_form"):
            col1, col2, col3 = st.columns(3)
            with col1:
                question_count = st.slider("🔢 Número de Exercícios", 1, 30, 5)
                grade_level = st.selectbox("🎓 Nível Escolar", ["Infantil", "Fundamental I", "Fundamental II", "Médio", "Superior"], index=2)
            with col2:
                difficulty = st.slider("📊 Dificuldade (0-100)", 0, 100, 40)
                question_type = st.selectbox("✍️ Tipo de Questão", ["Discursiva", "Múltipla Escolha"])
            with col3:
                include_images = st.checkbox("🖼️ Incluir imagens com Vertex AI", value=False)
                output_format = st.selectbox("💾 Formato de Saída", ["docx", "pdf", "txt"], help="Escolha DOCX ou PDF para incluir imagens.")
                image_style = st.text_input("🎨 Estilo da Imagem", "desenho vetorial simples", disabled=not include_images)
            special_need = st.selectbox("♿ Necessidade Específica (Opcional)", ["Nenhuma necessidade específica", "Síndrome de Down", "TEA", "Deficiência Intelectual", "TDAH"])
            submit_button = st.form_submit_button("🚀 Gerar Exercícios")
        
        if submit_button:
            if (generation_mode == "A partir de um Tema" and not theme_input.strip()) or \
               (generation_mode == "A partir de um Texto" and not user_text_input.strip()):
                st.error("Por favor, preencha o campo de tema ou texto.")
            else:
                with st.spinner("Gerando exercícios... Isso pode levar um momento."):
                    try:
                        payload = {'generationMode': 'theme' if generation_mode == "A partir de um Tema" else 'from_text', 'theme': theme_input, 'userText': user_text_input, 'questionCount': question_count, 'gradeLevel': grade_level, 'difficulty': difficulty, 'questionType': question_type, 'specialNeed': special_need, 'include_images': include_images, 'imageStyle': image_style, 'outputFormat': output_format}
                        file_stream, filename, mime = generate_exercises(payload, GEMINI_API_KEY)
                        st.success("✅ Exercícios gerados com sucesso!")
                        c1, c2 = st.columns(2)
                        with c1:
                           st.download_button(label="📥 Baixar Arquivo", data=file_stream, file_name=filename, mime=mime, use_container_width=True)
                        with c2:
                            save_stream = BytesIO(file_stream.getvalue())
                            if st.button("💾 Salvar em Meus Arquivos", use_container_width=True, key="save_exerc"):
                                save_file_to_user_storage(save_stream, filename)
                    except Exception as e:
                        st.error(f"❌ Erro ao gerar os exercícios: {e}")

    def page_otimizador_prompt():
        st.header("✨ Otimizador e Executor de Prompts com IA")
        st.markdown("Crie um prompt otimizado e use-o imediatamente para gerar texto ou imagens. Você pode incluir uma imagem como contexto.")

        # --- 1. CONFIGURAÇÃO E CHAVES DE API ---
        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')

        if not GEMINI_API_KEY:
            st.warning("A chave da API Gemini é necessária. Configure-a em 'Perfil e Configurações'.")
            return

        # --- 2. FUNÇÕES AUXILIARES ---
        def op_analisar_imagem_para_contexto(image_bytes):
            if not GCP_PROJECT_ID or not GCP_LOCATION:
                st.error("Project ID e Location do Google Cloud são necessários para analisar imagens.")
                return None
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                image_part = Part.from_data(image_bytes, mime_type="image/jpeg")
                prompt_instrucao = "Descreva de forma concisa e objetiva o conteúdo principal desta imagem. Esta descrição será usada como contexto para criar um novo prompt de IA. Foque nos elementos, estilo e composição."
                response = multimodal_model.generate_content([prompt_instrucao, image_part])
                return response.text
            except Exception as e:
                st.error(f"Erro ao analisar a imagem de contexto: {e}")
                return None

        def op_gerar_imagem_com_vertex(prompt):
            if not GCP_PROJECT_ID or not GCP_LOCATION:
                st.error("Project ID e Location do Google Cloud são necessários para gerar imagens.")
                return None
            try:
                if 'VERTEX_IMAGE_MODEL' not in globals():
                    st.error("O modelo de imagem Vertex AI não está definido.")
                    return None
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                modelo_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo_imagem.generate_images(prompt=prompt, number_of_images=1, aspect_ratio="1:1")
                if not resposta or not resposta[0]._image_bytes:
                    st.warning("A IA se recusou a gerar a imagem. Tente um prompt diferente.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Erro ao gerar a imagem via Vertex AI: {e}")
                return None

        # --- 3. INTERFACE DO USUÁRIO ---
        media_type = st.selectbox("1. Qual o objetivo final?", ("Gerar Texto", "Gerar Imagem"), key="op_media")
        
        st.markdown("---")
        st.subheader("2. Forneça o Contexto")

        uploaded_file = st.file_uploader("Envie uma imagem como inspiração ou contexto (Opcional)", type=["png", "jpg", "jpeg"])
        if uploaded_file:
            st.image(uploaded_file, caption="Imagem de Contexto", width=200)

        request_type = st.selectbox("Tipo de Requisição:", ("Geração de Conteúdo", "Resumo", "Tradução", "Análise de Sentimento", "Ideação", "Geração de Código", "Debug", "Refatoração", "Outro"), key="op_req")
        specific_details = st.text_area("Detalhes Específicos (tom, formato, público-alvo):", placeholder="Ex: Tom formal, formato de lista, para um público de desenvolvedores.", key="op_details")
        content = st.text_area("Requisitos (o que você deseja que a IA faça?):", placeholder="Ex: Escreva um artigo sobre os benefícios da IA para pequenas empresas.", height=150, key="op_content")
        
        st.markdown("---")
        st.subheader("3. Ajustes da IA")
        gemini_model = st.selectbox("Modelo Gemini para Otimização:", (GEMINI_FLASH_MODEL, GEMINI_PRO_MODEL), key="op_model")
        temperature = st.slider("Criatividade (Temperatura):", 0.0, 1.0, 0.7, 0.05, key="op_temp")

        if st.button("🚀 Otimizar Prompt e Gerar Conteúdo", use_container_width=True):
            if not content.strip():
                st.error("O campo 'Requisitos' é obrigatório.")
            else:
                image_context_description = ""
                if uploaded_file:
                    with st.spinner("Analisando imagem de contexto..."):
                        image_data = uploaded_file.getvalue()
                        image_context_description = op_analisar_imagem_para_contexto(image_data)
                        if not image_context_description:
                            st.error("Falha ao analisar a imagem. O processo continuará sem o contexto visual.")

                with st.spinner("Etapa 1: Otimizando seu prompt com IA..."):
                    prompt_for_gemini = (
                        f"Você é um engenheiro de prompts especialista. Sua tarefa é criar um prompt otimizado e detalhado em inglês para uma IA de {'imagem' if media_type == 'Gerar Imagem' else 'texto'}. "
                        f"Use as seguintes informações:\n"
                        f"- Objetivo: {request_type}\n"
                        f"- Detalhes (Tom, formato, público): {specific_details}\n"
                        f"- Requisitos do usuário: {content}\n"
                    )
                    if image_context_description:
                        prompt_for_gemini += f"- Contexto da imagem fornecida: {image_context_description}\n"
                    
                    prompt_for_gemini += "O prompt gerado deve ser direto, claro, rico em detalhes relevantes e pronto para ser usado pela IA de destino."

                    enhanced_prompt = get_gemini_response(prompt_for_gemini, gemini_model, temperature, GEMINI_API_KEY)

                if not enhanced_prompt or "Erro" in enhanced_prompt:
                    st.error(f"Não foi possível gerar o prompt otimizado. Resposta da API: {enhanced_prompt}")
                    return

                with st.expander("Ver Prompt Otimizado Gerado", expanded=True):
                    st.code(enhanced_prompt, language='text')

                st.markdown("---")
                st.subheader("✅ Resultado Final")

                with st.spinner(f"Etapa 2: Usando o prompt para gerar a {'imagem' if media_type == 'Gerar Imagem' else 'saída de texto'}..."):
                    if media_type == "Gerar Texto":
                        final_text_result = get_gemini_response(enhanced_prompt, GEMINI_PRO_MODEL, temperature, GEMINI_API_KEY)
                        if final_text_result:
                            st.markdown(final_text_result)
                        else:
                            st.error("Falha ao gerar o texto final.")

                    elif media_type == "Gerar Imagem":
                        if not GCP_PROJECT_ID or not GCP_LOCATION:
                            st.error("Para gerar imagens, o Project ID e a Location do Google Cloud devem ser configurados no seu perfil.")
                            return
                        generated_image_bytes = op_gerar_imagem_com_vertex(enhanced_prompt)
                        if generated_image_bytes:
                            st.image(generated_image_bytes, caption="Imagem gerada com base no prompt otimizado.", use_column_width=True)
                            st.download_button(
                                label="📥 Baixar Imagem (PNG)",
                                data=generated_image_bytes,
                                file_name="imagem_gerada_otimizador.png",
                                mime="image/png",
                                use_container_width=True
                            )
                        else:
                            st.error("Falha ao gerar a imagem final.")

    def page_analise_visual():
        st.header("👁️‍🗨️ Análise Visual para Geração de Conteúdo")
        st.markdown("Envie uma imagem (diagrama, UI, esquema, etc.) e a IA irá gerar código e um prompt baseado nela.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GEMINI_API_KEY or not GCP_PROJECT_ID or not GCP_LOCATION: return

        def av_analisar_imagem(image_bytes, prompt_usuario):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                image_part = Part.from_data(image_bytes, mime_type="image/png")

                instrucao_ia = f"""
                Analise a imagem fornecida. Com base em seu conteúdo visual, gere duas saídas distintas e claramente separadas:

                1.  **Geração de Código:** Crie um trecho de código relevante (Python, HTML, etc.) ou pseudocódigo que represente a lógica, estrutura ou os elementos visuais da imagem. Se for um diagrama, gere a classe correspondente. Se for um layout de UI, gere o HTML/CSS básico. Adicione um comentário no início do bloco de código indicando a linguagem.

                2.  **Geração de Prompt:** Crie um prompt de texto conciso e otimizado que descreva a imagem. Este prompt deve ser formatado para ser usado diretamente na ferramenta 'Otimizador de Prompt', capturando a essência da imagem para uma futura geração de conteúdo.

                Se o usuário fornecer uma instrução específica, leve-a em consideração: '{prompt_usuario}'
                """

                response = multimodal_model.generate_content([instrucao_ia, image_part])
                return response.text

            except google_exceptions.GoogleAPIError as e:
                st.error(f"Google API Error durante análise de imagem: {e.message}")
                return None
            except Exception as e:
                st.error(f"Erro ao analisar a imagem com o modelo multimodal: {e}")
                return None

        uploaded_file = st.file_uploader("Selecione uma imagem", type=["png", "jpg", "jpeg"])
        user_prompt = st.text_input("Instrução Específica (Opcional):", placeholder="Ex: 'Foco no formulário de login', 'Gere o código em Python'")

        if uploaded_file is not None:
            st.image(uploaded_file, caption="Imagem Carregada.", use_column_width=True)

            if st.button("Analisar Imagem e Gerar Conteúdo", use_container_width=True):
                with st.spinner("A IA está analisando a imagem..."):
                    try:
                        image_data = uploaded_file.getvalue()
                        resultado_analise = av_analisar_imagem(image_data, user_prompt)

                        if resultado_analise:
                            st.success("Análise concluída com sucesso!")
                            st.markdown("---")

                            partes = resultado_analise.split("Geração de Prompt:")
                            if len(partes) == 2:
                                st.subheader("Código Gerado a partir da Imagem:")
                                st.code(partes[0].replace("Geração de Código:", "").strip(), language='python') 
                                st.subheader("Prompt Otimizado para Descrever a Imagem:")
                                st.code(partes[1].strip(), language='text')
                            else:
                                st.subheader("Saída da IA (formato inesperado):")
                                st.code(resultado_analise, language='text')

                    except Exception as e:
                        st.error(f"Ocorreu um erro no processo: {e}")
    
    def page_criador_aplicativos():
        st.header("🏗️ Criador de Aplicativos (Scaffolding)")
        st.markdown("Gere a estrutura de arquivos e o código base para diferentes tipos de aplicações. Agora com prompts otimizados!")

        def ca_gerar_scaffold(project_name, app_type, extra_functions=None):
            templates = {
                "Streamlit App": {
                    "app.py": f"""
import streamlit as st

st.set_page_config(page_title="{project_name}")

def main():
    st.title("Bem-vindo ao {project_name}!")
    st.write("Este é um aplicativo Streamlit gerado pela Plataforma IA Evoluída.")

    name = st.text_input("Qual é o seu nome?")
    if name:
        st.write(f"Olá, {{name}}!")
{extra_functions or ""}
if __name__ == "__main__":
    main()
                    """,
                    "requirements.txt": "streamlit"
                },
                "API Flask": {
                    "app.py": f"""
from flask import Flask, jsonify

app = Flask(__name__)

@app.route('/')
def index():
    return jsonify({{"message": "Bem-vindo à API '{project_name}'!"}})

@app.route('/api/data')
def get_data():
    return jsonify({{"id": 1, "name": "Dado de Exemplo"}})
{extra_functions or ""}
if __name__ == '__main__':
    app.run(debug=True)
                    """,
                    "requirements.txt": "Flask"
                },
                "Script de Automação": {
                     "main.py": f"""
import os

def main():
    print("Iniciando o script de automação: {project_name}")
    # Exemplo: Listar arquivos no diretório atual
    files = os.listdir('.')
    print("Arquivos encontrados:")
    for f in files:
        print(f"- {{f}}")
    print("Script concluído.")
{extra_functions or ""}
if __name__ == "__main__":
    main()
                    """,
                    "README.md": f"# {project_name}\\n\\nEste é um script de automação gerado pela Plataforma IA Evoluída."
                }
            }
            return templates.get(app_type, {})

        st.subheader("1. Defina os Detalhes do Projeto")
        project_name_input = st.text_input("Nome do Projeto:", "MeuNovoApp")
        app_type_select = st.selectbox(
            "Selecione o Tipo de Aplicativo:",
            ("Streamlit App", "API Flask", "Script de Automação")
        )

        st.subheader("2. Adicione Funções Extras com IA")
        extra_func_description = st.text_area(
            "Descreva funções extras para adicionar ao aplicativo (ex: exportar dados, autenticação, gráficos):",
            placeholder="Ex: Adicionar função para exportar dados em CSV."
        )

        GEMINI_API_KEY = get_api_key('gemini_key')
        gemini_model = GEMINI_FLASH_MODEL
        temperature = 0.7

        extra_functions_code = ""
        if extra_func_description.strip():
            if not GEMINI_API_KEY:
                st.warning("Chave Gemini API não configurada. Funções extras via IA não serão geradas.")
            else:
                with st.spinner("Otimizando prompt e gerando funções extras..."):
                    prompt_for_gemini = (
                        f"Gere funções Python para {app_type_select} conforme a descrição: {extra_func_description}. "
                        "O código deve ser pronto para uso e fácil de integrar ao template padrão."
                    )
                    enhanced_code = get_gemini_response(prompt_for_gemini, gemini_model, temperature, GEMINI_API_KEY)
                    if enhanced_code:
                        extra_functions_code = "\n" + enhanced_code

        if st.button("Gerar Estrutura do Aplicativo", use_container_width=True):
            if not project_name_input.strip():
                st.error("Por favor, forneça um nome para o projeto.")
            else:
                with st.spinner("Gerando arquivos do projeto..."):
                    scaffold_files = ca_gerar_scaffold(project_name_input, app_type_select, extra_functions=extra_functions_code)
                    if scaffold_files:
                        st.success(f"Estrutura para '{project_name_input}' gerada!")
                        st.balloons()
                        for filename, content in scaffold_files.items():
                            with st.expander(f"📄 Arquivo: {filename}"):
                                st.code(content, language='python' if filename.endswith('.py') else 'text')
                    else:
                        st.error("Tipo de aplicativo não encontrado.")

    def page_fabrica_spritesheets():
        st.header("👹 Fábrica de Spritesheets 2D - Geração por Template Avançado")
        st.markdown("Utiliza um template de prompt de alta qualidade para gerar a spritesheet completa de uma só vez.")

        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GCP_PROJECT_ID or not GCP_LOCATION: return

        def fs_gerar_prompt_template_avancado(data, lista_frames, estilo_arte):
            style_text = ("2D pixel art spritesheet" if estilo_arte == "Pixel Art" else "2D cartoon spritesheet")
            character_text = (f"a powerful '{data['funcao']}' character, wearing a {data['cor']}. "
                              f"Specific details: {data['detalhes']}. Wielding a {data['arma']}.")
            animation_text = ", ".join(lista_frames)
            prompt_final = (f"{style_text}, of {character_text}. "
                            f"The spritesheet must include a complete set of animation frames: {animation_text}. "
                            "Present the sprites in a clear, organized grid layout with a transparent background. "
                            "The character design must be strictly consistent across all poses. "
                            "Use a side-view perspective, optimized for smooth game animation. High resolution image.")
            return prompt_final

        def fs_gerar_imagem_google(prompt):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                st.info(f"Enviando o prompt para a IA...")
                modelo = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo.generate_images(prompt=prompt, number_of_images=1, aspect_ratio="16:9")
                if not resposta or not resposta[0]._image_bytes:
                    st.warning("A IA se recusou a gerar a imagem. Tente uma descrição diferente.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Erro ao gerar a imagem via Vertex AI: {e}")
                return None

        st.subheader("1. Defina os Detalhes do Personagem")
        estilo_arte_selecionado = st.selectbox("Estilo de Arte", ["Pixel Art", "Cartoon"], key="fs_style_select")
        col1, col2 = st.columns(2)
        with col1:
            funcao = st.text_input("Função ou Classe", "Mago Arcano", key="fs_funcao")
            arma = st.text_input("Arma/Equipamento", "cajado de madeira com orbe brilhante", key="fs_arma")
        with col2:
            cor = st.text_input("Cores Dominantes", "manto azul escuro e túnica roxa", key="fs_cor")
        detalhes = st.text_area("Detalhes Visuais", "ancião com longa barba branca e expressão sábia", key="fs_detalhes", height=100)

        st.subheader("2. Selecione as Animações")
        frames_disponiveis = ["Idle", "Walk", "Run", "Jump (take-off and landing)", "Basic Melee Attack (staff swing)", "Magic Spell Cast (fireball)", "Hurt", "Block", "Death animation"]
        frames_selecionados = st.multiselect("Selecione os frames para a spritesheet", options=frames_disponiveis, default=["Idle", "Walk", "Run", "Magic Spell Cast (fireball)"])

        if st.button("👹 Gerar Spritesheet", use_container_width=True):
            if not frames_selecionados:
                st.warning("Selecione pelo menos um frame de animação.")
            else:
                try:
                    with st.spinner("Construindo prompt e gerando a imagem..."):
                        dados_personagem = {"funcao": funcao, "detalhes": detalhes, "arma": arma, "cor": cor}
                        prompt_completo = fs_gerar_prompt_template_avancado(dados_personagem, frames_selecionados, estilo_arte_selecionado)
                        imagem_gerada_bytes = fs_gerar_imagem_google(prompt_completo)

                        if imagem_gerada_bytes:
                            st.success("Spritesheet gerada com sucesso!")
                            st.image(imagem_gerada_bytes, caption="Resultado da geração com o template.", use_container_width=True)
                            st.download_button("📥 Baixar Spritesheet (PNG)", data=imagem_gerada_bytes, file_name=f"{funcao.replace(' ', '_').lower()}_spritesheet.png", mime="image/png", use_container_width=True)
                except Exception as e:
                    st.error(f"Erro ao gerar imagens com Vertex AI: {e}")
    
    def page_analise_logs():
        st.header("📊 Análise de Logs e Eventos com IA")
        st.markdown("Cole o conteúdo do seu log ou faça upload de uma imagem (ex: screenshot de erro) para a IA analisar e fornecer insights.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')

        if not GEMINI_API_KEY or not GCP_PROJECT_ID or not GCP_LOCATION:
            st.warning("Configure suas chaves de API para Gemini, Google Cloud Project ID e Location na aba 'Perfil e Configurações' para usar esta ferramenta.")
            return

        def al_analisar_imagem_multimodal(image_bytes, prompt_usuario):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                image_part = Part.from_data(image_bytes, mime_type="image/png")

                instrucao_ia = f"""
                Analise esta imagem que faz parte de um log ou representa um evento do sistema.
                Extraia qualquer texto visível (OCR), identifique diagramas, gráficos ou estruturas visuais.
                Descreva o conteúdo visual de forma concisa e relevante para a análise de logs.
                Se for um diagrama de fluxo, descreva o fluxo. Se for um erro, transcreva-o.
                Sua análise deve ser útil para complementar a compreensão do log textual.
                Instrução específica do usuário: '{prompt_usuario}'
                """
                response = multimodal_model.generate_content([instrucao_ia, image_part])
                return response.text
            except google_exceptions.GoogleAPIError as e:
                st.error(f"Erro na API do Google ao analisar a imagem: {e.message}")
                return f"Erro na análise visual da imagem: {e.message}"
            except Exception as e:
                st.error(f"Ocorreu um erro inesperado ao analisar a imagem: {e}")
                return f"Erro na análise visual da imagem: {e}"

        with st.form(key="log_analysis_form"):
            st.subheader("1. Forneça o Conteúdo do Log")
            log_content = st.text_area("📜 Cole o conteúdo do log aqui:", height=250,
                                       placeholder="Ex: [2023-10-26 14:30:15] ERROR: Falha na conexão com o banco de dados. Erro: Timeout, host=db.example.com")

            uploaded_image_log = st.file_uploader("🖼️ Ou faça upload de uma imagem de log/evento (opcional):", type=["png", "jpg", "jpeg"])
            image_analysis_instruction = st.text_input("Instrução para análise da imagem (opcional):",
                                                        placeholder="Ex: 'Foque no diagrama de rede', 'Transcreva o erro na tela'")

            st.subheader("2. Detalhes Adicionais (Opcional)")
            col1, col2, col3 = st.columns(3)
            with col1:
                log_type = st.selectbox("Tipo de Log:", ["Geral", "Sistema", "Aplicação", "Segurança", "Rede", "Banco de Dados", "Outro"])
            with col2:
                data_inicio = st.date_input("Data de Início:", value=None, key="log_start_date")
            with col3:
                data_fim = st.date_input("Data de Fim:", value=None, key="log_end_date")

            col4, col5 = st.columns(2)
            with col4:
                gravidade = st.selectbox("Gravidade:", ["Qualquer", "DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"])
            with col5:
                data_source = st.text_input("Fonte de Dados (Sistema/Serviço):", placeholder="Ex: Apache, Nginx, Kubernetes, Servidor DB")

            user_instruction = st.text_area("Instrução específica para a IA (como você quer a análise?):", height=100,
                                            placeholder="Ex: 'Foque em anomalias de desempenho', 'Identifique tentativas de acesso não autorizado'")

            submit_button = st.form_submit_button("🚀 Analisar Log com IA")

        if submit_button:
            if not log_content and not uploaded_image_log:
                st.error("Por favor, cole o conteúdo do log ou faça upload de uma imagem.")
            else:
                with st.spinner("Analisando log e eventos..."):
                    try:
                        image_analysis_result = ""
                        if uploaded_image_log:
                            st.info("Processando imagem de log...")
                            image_bytes = uploaded_image_log.getvalue()
                            image_analysis_result = al_analisar_imagem_multimodal(image_bytes, image_analysis_instruction)
                            if image_analysis_result:
                                st.success("Análise visual concluída.")
                            else:
                                st.warning("Não foi possível obter uma análise visual da imagem.")


                        prompt_parts = []
                        prompt_parts.append("Analise o seguinte log e eventos para identificar padrões, anomalias, problemas potenciais e sugerir melhorias. Forneça um relatório detalhado.")
                        prompt_parts.append(f"\n--- Detalhes do Log ---")
                        prompt_parts.append(f"Tipo de Log: {log_type}")
                        if data_inicio: prompt_parts.append(f"Data de Início: {data_inicio}")
                        if data_fim: prompt_parts.append(f"Data de Fim: {data_fim}")
                        prompt_parts.append(f"Gravidade: {gravidade}")
                        if data_source: prompt_parts.append(f"Fonte de Dados: {data_source}")

                        if log_content:
                            prompt_parts.append(f"\n--- Conteúdo Textual do Log ---\n{log_content}")

                        if image_analysis_result:
                            prompt_parts.append(f"\n--- Análise Visual da Imagem ---\n{image_analysis_result}")

                        if user_instruction:
                            prompt_parts.append(f"\n--- Instruções Específicas do Usuário ---\n{user_instruction}")

                        prompt_parts.append(f"\n--- Estrutura do Relatório Requerida ---")
                        prompt_parts.append("O relatório deve ser estruturado nas seguintes seções: ")
                        prompt_parts.append("1. **Sumário Executivo:** Breve resumo das descobertas mais importantes.")
                        prompt_parts.append("2. **Eventos Chave e Linha do Tempo:** Descrição dos eventos mais críticos ou relevantes, com marcadores de tempo se disponíveis no log.")
                        prompt_parts.append("3. **Padrões Identificados:** Quaisquer repetições, tendências ou comportamentos normais/anormais encontrados.")
                        prompt_parts.append("4. **Anomalias e Problemas:** Detalhes sobre erros, exceções, falhas ou comportamentos inesperados.")
                        prompt_parts.append("5. **Impacto Potencial:** Explicação de como os problemas identificados podem afetar o sistema ou os usuários.")
                        prompt_parts.append("6. **Recomendações e Próximos Passos:** Sugestões acionáveis para resolver problemas, otimizar ou prevenir ocorrências futuras.")
                        prompt_parts.append("Use formatação Markdown para clareza.")

                        full_prompt = "\n".join(prompt_parts)

                        analysis_result = get_gemini_response(full_prompt, GEMINI_PRO_MODEL, temperature=0.2, api_key=GEMINI_API_KEY)

                        if analysis_result:
                            st.success("✅ Análise de Log Concluída!")
                            st.markdown(analysis_result)
                        else:
                            st.error("❌ Não foi possível gerar a análise do log. Tente novamente ou verifique suas chaves de API.")

                    except Exception as e:
                        st.error(f"❌ Ocorreu um erro inesperado durante a análise: {e}")

    # --- INÍCIO DA NOVA FERRAMENTA ---
    def page_criador_pop_retrogaming():
        st.header("🛠️ Gerador de Documentação POP para Retrogaming")
        st.markdown("Crie guias de suporte e procedimentos operacionais padrão (POP) para sistemas de retrogaming com a ajuda da IA.")

        # --- Funções de Apoio Específicas para esta Página ---
        def pop_analisar_imagem_contexto(image_bytes, gcp_project_id, gcp_location):
            """Usa o modelo multimodal para descrever uma imagem para o contexto do POP."""
            try:
                vertexai.init(project=gcp_project_id, location=gcp_location)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                image_part = Part.from_data(image_bytes, mime_type="image/jpeg")
                prompt_instrucao = """
                Analise a imagem fornecida, que provavelmente é um screenshot de um sistema de emulação (como Batocera, RetroArch), um erro de software, ou uma foto de hardware.
                Descreva de forma objetiva e técnica o que você vê.
                - Se for um erro, transcreva a mensagem de erro.
                - Se for uma tela de configuração, descreva as opções visíveis.
                - Se for hardware, descreva os componentes e conexões.
                Esta descrição será usada para criar um documento de suporte técnico.
                """
                response = multimodal_model.generate_content([prompt_instrucao, image_part])
                return response.text
            except Exception as e:
                st.error(f"Erro ao analisar a imagem de contexto: {e}")
                return None

        def pop_criar_docx(content, image_bytes=None):
            """Cria um arquivo DOCX em memória com o conteúdo e a imagem."""
            document = Document()
            document.add_heading('Documentação POP - Retrogaming', 0)
            
            if image_bytes:
                try:
                    image_stream = BytesIO(image_bytes)
                    document.add_picture(image_stream, width=document.sections[0].page_width * 0.7)
                    document.add_paragraph("Imagem de Referência", style='Caption')
                except Exception as e:
                    st.warning(f"Não foi possível adicionar a imagem ao DOCX: {e}")

            document.add_paragraph(content)
            
            file_stream = BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            return file_stream

        def pop_criar_pdf(content, image_bytes=None):
            """Cria um arquivo PDF em memória com o conteúdo e a imagem."""
            file_stream = BytesIO()
            doc = SimpleDocTemplate(file_stream, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            story.append(Paragraph("Documentação POP - Retrogaming", styles['h1']))
            story.append(Spacer(1, 12))

            if image_bytes:
                try:
                    image_stream = BytesIO(image_bytes)
                    img = Image(image_stream, width=400, height=300)
                    story.append(img)
                    story.append(Spacer(1, 12))
                except Exception as e:
                    st.warning(f"Não foi possível adicionar a imagem ao PDF: {e}")
            
            content_html = content.replace('\n', '<br/>')
            story.append(Paragraph(content_html, styles['BodyText']))

            doc.build(story)
            file_stream.seek(0)
            return file_stream

        # --- Interface do Usuário ---
        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')

        if not all([GEMINI_API_KEY, GCP_PROJECT_ID, GCP_LOCATION]):
            st.warning("Configure as chaves Gemini API, Google Cloud Project ID e Location em 'Perfil e Configurações' para usar esta ferramenta.")
            return

        with st.form(key="pop_generator_form"):
            st.subheader("1. Detalhes da Documentação")
            
            pop_objective = st.text_input("Qual o objetivo deste POP?", placeholder="Ex: Corrigir erro de tela preta no emulador de PS2")
            
            pop_system = st.selectbox(
                "Selecione o Sistema/Software:",
                ("Sistema (Geral, Windows, Linux)", "Batocera", "RetroPie", "Lakka", "MAME", "Hyperspin", "CoinOps", "Multi-jogos (Build Custom)", "Emulador Específico", "Hardware (Controles, Gabinete)")
            )
            
            pop_details = st.text_area("Descreva o problema ou o procedimento:", height=200, placeholder="Ex: Ao iniciar qualquer jogo de PS2 no Batocera, a tela fica preta, mas o som continua. Isso acontece na versão 35 com uma GPU Nvidia.")

            st.subheader("2. Recursos Adicionais")
            uploaded_image = st.file_uploader("Envie uma imagem de referência (opcional)", type=["png", "jpg", "jpeg"])

            st.subheader("3. Configurações de Geração")
            col1, col2 = st.columns(2)
            with col1:
                selected_model = st.selectbox("Escolha o modelo de IA:", (GEMINI_FLASH_MODEL, GEMINI_PRO_MODEL), help="Flash é mais rápido, Pro é mais detalhado.")
            with col2:
                output_format = st.selectbox("Formato de Saída:", ("docx", "pdf"))

            submit_button = st.form_submit_button("🚀 Gerar Documentação POP")

        if submit_button:
            if not pop_objective or not pop_details:
                st.error("Por favor, preencha os campos 'Objetivo' e 'Descrição'.")
            else:
                with st.spinner("Analisando informações e gerando documentação..."):
                    try:
                        image_context = ""
                        image_bytes = None
                        if uploaded_image:
                            st.info("Analisando imagem enviada...")
                            image_bytes = uploaded_image.getvalue()
                            image_context = pop_analisar_imagem_contexto(image_bytes, GCP_PROJECT_ID, GCP_LOCATION)
                            if image_context:
                                st.success("Análise da imagem concluída.")
                            else:
                                st.warning("Não foi possível analisar a imagem. O POP será gerado sem este contexto.")
                        
                        prompt_parts = [
                            "Você é um especialista em sistemas de retrogaming e um escritor técnico. Sua tarefa é criar um Procedimento Operacional Padrão (POP) claro, conciso e fácil de seguir.",
                            "O documento deve ser estruturado profissionalmente com as seguintes seções, usando Markdown para formatação:",
                            "- **Título:** Um título claro para o procedimento.",
                            "- **Objetivo:** O que este procedimento visa resolver ou alcançar.",
                            "- **Sistema Aplicável:** O sistema ou software ao qual se aplica.",
                            "- **Ferramentas/Requisitos:** Qualquer software ou hardware necessário.",
                            "- **Procedimento Passo a Passo:** Instruções numeradas e detalhadas.",
                            "- **Verificação:** Como confirmar que o procedimento foi bem-sucedido.",
                            "- **Notas e Solução de Problemas (Opcional):** Dicas adicionais ou problemas comuns.",
                            "\n--- DADOS FORNECIDOS PELO USUÁRIO ---",
                            f"**Objetivo do POP:** {pop_objective}",
                            f"**Sistema:** {pop_system}",
                            f"**Descrição do Problema/Procedimento:**\n{pop_details}"
                        ]
                        if image_context:
                            prompt_parts.append(f"\n**Contexto da Imagem Anexada (Analisada pela IA):**\n{image_context}")
                        
                        full_prompt = "\n".join(prompt_parts)

                        st.info("Gerando o texto do POP com a IA...")
                        generated_content = get_gemini_response(full_prompt, selected_model, temperature=0.5, api_key=GEMINI_API_KEY)
                        
                        if generated_content and "Erro:" not in generated_content:
                            st.success("✅ Documentação gerada com sucesso!")
                            st.session_state['generated_pop_content'] = generated_content
                            st.session_state['generated_pop_image_bytes'] = image_bytes
                            st.session_state['output_format'] = output_format
                        else:
                            st.error(f"❌ Falha ao gerar a documentação: {generated_content}")

                    except Exception as e:
                        st.error(f"❌ Ocorreu um erro inesperado: {e}")

        if 'generated_pop_content' in st.session_state:
            st.markdown("---")
            st.subheader("Resultado Gerado")
            
            st.markdown(st.session_state['generated_pop_content'])
            
            file_stream = None
            if st.session_state['output_format'] == 'docx':
                file_stream = pop_criar_docx(st.session_state['generated_pop_content'], st.session_state.get('generated_pop_image_bytes'))
                file_name = "documentacao_pop.docx"
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                file_stream = pop_criar_pdf(st.session_state['generated_pop_content'], st.session_state.get('generated_pop_image_bytes'))
                file_name = "documentacao_pop.pdf"
                mime_type = "application/pdf"
            
            if file_stream:
                st.download_button(
                    label=f"📥 Baixar POP em .{st.session_state['output_format'].upper()}",
                    data=file_stream,
                    file_name=file_name,
                    mime=mime_type,
                    use_container_width=True
                )
    # --- FIM DA NOVA FERRAMENTA ---

    def page_espelho_da_mente():
        st.header("✨ Espelho da Mente Dinâmico")
        st.markdown("Transforme pensamentos e sentimentos complexos em arte simbólica.")

        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GCP_PROJECT_ID or not GCP_LOCATION: return

        def em_criar_prompt_avancado_com_ia(dados_usuario):
            GEMINI_API_KEY = get_api_key('gemini_key')
            if not GEMINI_API_KEY:
                return None
            modelo_texto = GenerativeModel(GEMINI_FLASH_MODEL)
            prompt_base = f"""
            Sua tarefa é atuar como um engenheiro de prompt especialista em geração de imagens artísticas.
            Converta a descrição do usuário em um prompt em inglês, detalhado e evocativo para um modelo de IA de imagem.
            O objetivo é criar uma metáfora visual do sentimento do usuário.

            **Diretrizes:**
            1. **Estilo e Atmosfera:** O estilo principal é '{dados_usuario['estilo']}'. Descreva a atmosfera usando luz, sombra e a paleta de cores '{dados_usuario['cores']}' para intensificar a emoção.
            2. **Composição:** Cenário: '{dados_usuario['cenario']}'. Ponto focal: '{dados_usuario['objeto']}', que simboliza a experiência.
            3. **Simbolismo:** Incorpore o sentimento '{dados_usuario['pensamento']}' na cena de forma sutil.
            4. **Detalhes Adicionais:** Inclua: '{dados_usuario['detalhes']}'.
            5. **Restrições:** Evite: '{dados_usuario['evitar']}'.

            Gere o prompt final para a IA de imagem.
            """
            try:
                response = modelo_texto.generate_content([prompt_base])
                return response.text.strip()
            except google_exceptions.GoogleAPIError as e:
                st.error(f"Erro na API do Google ao criar prompt para imagem: {e.message}")
                return None
            except Exception as e:
                st.error(f"Ocorreu um erro inesperado ao criar o prompt: {e}")
                return None

        def em_gerar_imagem_google(prompt):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                modelo_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo_imagem.generate_images(prompt=prompt, number_of_images=1, aspect_ratio="1:1")
                if not resposta or not resposta[0]._image_bytes:
                    st.warning("A IA recusou-se a gerar a imagem. Tente uma descrição diferente.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Erro ao gerar a imagem via Vertex AI: {e}")
                return None

        st.subheader("🎨 Construa sua Visão")
        pensamento_usuario = st.text_area("O que você está pensando ou sentindo?", "Nostalgia e esperança por um futuro incerto.", height=100, key="em_thought")
        cenario = st.text_input("Se esse sentimento fosse um lugar, qual seria?", "Uma estação de trem antiga e vazia ao amanhecer.", key="em_place")
        objeto = st.text_input("Qual objeto simboliza melhor este momento?", "Um único broto verde crescendo entre os trilhos.", key="em_object")
        estilo = st.selectbox("Em qual estilo de arte?", ("Surrealismo", "Impressionismo", "Arte Conceitual", "Abstrato", "Aquarela", "Fantasia Sombria", "Cyberpunk"), key="em_style")
        cores = st.text_input("Quais as cores ou paleta?", "Tons de sépia, cinza, com um ponto de luz dourada.", key="em_colors")
        detalhes = st.text_input("Algum detalhe adicional?", "Névoa baixa no chão, raios de sol por uma janela quebrada.", key="em_details_add")
        evitar = st.text_input("O que NÃO incluir na imagem?", "Pessoas, animais, relógios.", key="em_avoid")

        if st.button("✨ Gerar Imagem da Emoção", use_container_width=True):
            if not all([pensamento_usuario, cenario, objeto]):
                st.warning("Por favor, preencha os campos de sentimento, lugar e objeto.")
            else:
                try:
                    with st.spinner("Etapa 1: Interpretando seus sentimentos com IA..."):
                        dados_para_prompt = {"pensamento": pensamento_usuario, "cenario": cenario, "objeto": objeto, "estilo": estilo, "cores": cores, "detalhes": detalhes, "evitar": evitar}
                        prompt_final = em_criar_prompt_avancado_com_ia(dados_para_prompt)

                    if prompt_final:
                        with st.expander("Ver o prompt técnico gerado"):
                            st.write(prompt_final)

                        with st.spinner("Etapa 2: Pintando sua visão..."):
                            imagem_bytes = em_gerar_imagem_google(prompt_final)

                        if imagem_bytes:
                            st.success("Sua imagem foi criada!")
                            st.image(imagem_bytes, caption="Uma reflexão visual do seu sentimento.", use_column_width=True)
                            st.download_button("📥 Baixar Imagem (PNG)", data=imagem_bytes, file_name="espelho_da_mente.png", mime="image/png", use_container_width=True)
                    else:
                        st.error("Não foi possível gerar o prompt para a imagem.")
                except Exception as e:
                    st.error(f"Erro ao gerar a imagem: {e}")

    def page_pesquisa_avancada():
            st.header("🔎 Ferramenta de Pesquisa Avançada (Google Dorks)")
            st.markdown("Realize buscas precisas no Google usando operadores dork ou faça uma busca reversa a partir de uma imagem.")

            # --- CHAVES DE API E FUNÇÕES DE APOIO ---
            GSEARCH_KEY = get_api_key('gsearch_key')
            GSEARCH_CX = get_api_key('gsearch_cx')
            GEMINI_API_KEY = get_api_key('gemini_key')
            GCP_PROJECT_ID = get_api_key('gcp_project_id')
            GCP_LOCATION = get_api_key('gcp_location')

            if not GSEARCH_KEY or not GSEARCH_CX or not GEMINI_API_KEY:
                st.warning("⚠️ Configure suas chaves para Google Search e Gemini na aba 'Perfil e Configurações' para usar esta ferramenta.")
                return

            def pa_google_search(query, api_key, cx):
                """Função genérica para realizar a busca."""
                try:
                    service = build("customsearch", "v1", developerKey=api_key)
                    res = service.cse().list(q=query, cx=cx, num=10).execute()
                    return res.get("items", [])
                except Exception as e:
                    st.error(f"Erro ao realizar a busca no Google: {e}")
                    return []
            
            def pa_analisar_imagem_para_busca(image_bytes):
                """Usa o Gemini para descrever a imagem e gerar um termo de busca."""
                if not GCP_PROJECT_ID or not GCP_LOCATION:
                    st.error("Project ID e Location do Google Cloud são necessários para busca por imagem.")
                    return None
                try:
                    vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                    multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                    image_part = Part.from_data(image_bytes, mime_type="image/jpeg")
                    prompt_instrucao = "Descreva esta imagem de forma concisa para uma busca na web. Foque nos objetos, pessoas, locais ou texto visível."
                    response = multimodal_model.generate_content([prompt_instrucao, image_part])
                    return response.text
                except Exception as e:
                    st.error(f"Erro ao analisar a imagem para busca: {e}")
                    return None

            # --- INTERFACE COM ABAS ---
            tab1, tab2 = st.tabs(["Pesquisa por Texto & Dorks", "Pesquisa Reversa por Imagem"])

            with tab1:
                st.subheader("Busca por Texto com Operadores Dork")
                st.info("Preencha os campos para construir sua busca. Não é preciso digitar os operadores como 'site:' ou 'filetype:'.")

                # Dorks pré-configurados para facilitar o uso
                dork_templates = {
                    "Nenhum (busca padrão)": "",
                    "Encontrar Arquivos (PDF/DOC/XLS)": 'ext:(doc | pdf | xls | txt)',
                    "Encontrar Diretórios Abertos (Música/Vídeo)": 'intitle:"index.of" "parent directory" (mp4|mp3|avi|flac)',
                    "Encontrar Arquivos de Configuração": 'filetype:config OR filetype:ini OR filetype:log',
                    "Encontrar Páginas de Login": 'inurl:login OR inurl:signin OR intitle:"login"',
                    "Encontrar Documentos Confidenciais": '(intext:confidential salary | intext:"budget approved")',
                    "Pesquisar em sites .gov e .edu": 'site:*.gov OR site:*.edu'
                }

                with st.form("dork_form"):
                    search_terms = st.text_input("📝 **Termos Principais da Pesquisa**", placeholder="Ex: Relatório financeiro, Nina Simone, vulnerabilidade XSS")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        site_filter = st.text_input("🌐 **Restringir ao site (ex: google.com)**", key="site_filter")
                        url_terms = st.text_input("🔗 **Termos que devem estar na URL**", key="url_terms", placeholder="Ex: admin, dashboard")
                    with col2:
                        file_types = st.multiselect("📄 **Tipo de Arquivo Específico**", 
                                                    options=["pdf", "doc", "docx", "xls", "xlsx", "ppt", "pptx", "txt", "log", "sql"],
                                                    key="file_types")
                    
                    template_choice = st.selectbox("🚀 **Usar um Dork Pronto (Template)**", options=list(dork_templates.keys()))

                    submit_dork_search = st.form_submit_button("🔎 Pesquisar com Dorks", use_container_width=True)

                if submit_dork_search and search_terms:
                    with st.spinner("Construindo dork e pesquisando..."):
                        query_parts = [search_terms]
                        if site_filter:
                            query_parts.append(f'site:{site_filter}')
                        if url_terms:
                            query_parts.append(f'inurl:{url_terms}')
                        if file_types:
                            filetype_query = " OR ".join([f'filetype:{ft}' for ft in file_types])
                            query_parts.append(f'({filetype_query})')
                        if template_choice != "Nenhum (busca padrão)":
                            query_parts.append(dork_templates[template_choice])

                        final_dork = " ".join(query_parts)
                        st.success(f"**Dork Executado:** `{final_dork}`")
                        st.session_state['search_results'] = pa_google_search(final_dork, GSEARCH_KEY, GSEARCH_CX)

            with tab2:
                st.subheader("Busca a partir de uma Imagem")
                st.info("Envie uma imagem e a IA irá descrevê-la para buscar imagens ou contextos semelhantes na web.")
                
                uploaded_image = st.file_uploader("Selecione uma imagem para a busca", type=["png", "jpg", "jpeg"])
                
                if uploaded_image:
                    st.image(uploaded_image, caption="Imagem para análise", width=250)
                
                extra_context = st.text_input("Adicione contexto à sua busca (opcional):", placeholder="Ex: Onde foi tirada essa foto?, Que tipo de planta é essa?")

                if st.button("📸 Pesquisar com Imagem", use_container_width=True):
                    if uploaded_image:
                        with st.spinner("Analisando a imagem com IA e realizando a busca..."):
                            image_bytes = uploaded_image.getvalue()
                            image_description = pa_analisar_imagem_para_busca(image_bytes)
                            
                            if image_description:
                                st.success(f"**IA descreveu a imagem como:** '{image_description}'")
                                search_query = f"{image_description} {extra_context}".strip()
                                st.session_state['search_results'] = pa_google_search(search_query, GSEARCH_KEY, GSEARCH_CX)
                            else:
                                st.error("Não foi possível analisar a imagem. A busca não foi realizada.")
                    else:
                        st.warning("Por favor, envie uma imagem antes de pesquisar.")

            # --- EXIBIÇÃO DOS RESULTADOS (COMUM A AMBAS AS ABAS) ---
            if 'search_results' in st.session_state:
                results = st.session_state.get('search_results', [])
                if results:
                    st.markdown("---")
                    st.subheader(f"✅ Resultados Encontrados ({len(results)})")
                    for r in results:
                        st.markdown(f"#### [{r.get('title', 'Sem título')}]({r.get('link', '#')})")
                        st.caption(f"🌐 {r.get('displayLink', 'N/A')}")
                        st.write(r.get('snippet', 'Sem descrição.'))
                        st.markdown("---")
                else:
                    st.warning("Nenhum resultado encontrado para sua busca.")
   
    def page_narrador_ia():
        st.header("🎙️ Narrador IA - Texto para Áudio")
        st.markdown("Transforme texto em narrações ou extraia, traduza e narre o conteúdo de vídeos do YouTube.")

        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GEMINI_API_KEY = get_api_key('gemini_key')
        
        if not GCP_PROJECT_ID or not GEMINI_API_KEY:
            st.warning("O ID do Projeto Google Cloud e a Chave Gemini são necessários. Configure no seu perfil.")
            return

        def ni_get_youtube_video_id(url):
            if not url: return None
            regex = r"(?:https?:\/\/)?(?:www\.)?(?:youtube\.com\/(?:[^\/\n\s]+\/\S+\/|(?:v|e(?:mbed)?)\/|\S*?[?&]v=)|youtu\.be\/)([a-zA-Z0-9_-]{11})"
            match = re.search(regex, url)
            return match.group(1) if match else None
        
        def ni_sintetizar_fala(texto, language_code, voice_name, speaking_rate):
            try:
                texto = texto.strip()
                max_bytes = 4500  
                partes = []
                parte_atual = ""

                for linha in texto.split(". "):
                    if len((parte_atual + linha).encode("utf-8")) < max_bytes:
                        parte_atual += (linha + ". ")
                    else:
                        partes.append(parte_atual.strip())
                        parte_atual = linha + ". "
                if parte_atual:
                    partes.append(parte_atual.strip())

                client = texttospeech.TextToSpeechClient()
                audio_total = b""

                for parte in partes:
                    synthesis_input = texttospeech.SynthesisInput(text=parte)
                    voice = texttospeech.VoiceSelectionParams(
                        language_code=language_code,
                        name=voice_name
                    )
                    audio_config = texttospeech.AudioConfig(
                        audio_encoding=texttospeech.AudioEncoding.MP3,
                        speaking_rate=speaking_rate
                    )
                    response = client.synthesize_speech(
                        input=synthesis_input,
                        voice=voice,
                        audio_config=audio_config
                    )
                    audio_total += response.audio_content

                return {"success": True, "data": audio_total}

            except Exception as e:
                logging.error(f"Erro ao converter texto em áudio: {e}")
                return {"success": False, "error": f"Erro na API de áudio: {e}"}

        def ni_get_and_translate_transcript(video_url, gemini_api_key):
            try:
                video_id = ni_get_youtube_video_id(video_url)
                if not video_id:
                    return {"success": False, "error": "URL do YouTube inválida ou formato não reconhecido."}
                
                st.info(f"Buscando transcrição para o vídeo ID: {video_id}...")
                
                transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
                transcript = transcript_list.find_transcript(['en', 'es', 'pt', 'de', 'fr']) # Prioritize languages
                transcript_data = transcript.fetch()
                
                original_transcript = ' '.join([item['text'] for item in transcript_data])
                
                st.info("Tradução em andamento com a IA...")
                translation_prompt = (f"Traduza o seguinte texto para o Português do Brasil. Mantenha a estrutura e o significado o mais fiel possível. "
                                      f"Se o texto já estiver em um bom português, apenas o retorne. Não adicione nenhuma introdução ou comentário seu, apenas retorne o texto traduzido.\n\n"
                                      f"TEXTO ORIGINAL:\n---\n{original_transcript}")
                
                translated_text = get_gemini_response(translation_prompt, GEMINI_PRO_MODEL, 0.2, gemini_api_key)
                
                if "Erro:" in translated_text:
                    return {"success": False, "error": f"Falha na tradução: {translated_text}"}
                
                return {"success": True, "data": translated_text}

            except TranscriptsDisabled:
                return {"success": False, "error": f"As legendas estão desativadas para o vídeo com ID '{video_id}'."}
            except NoTranscriptFound:
                return {"success": False, "error": f"Não foi encontrada nenhuma legenda para o vídeo com ID '{video_id}'."}
            except VideoUnavailable:
                 return {"success": False, "error": f"O vídeo com ID '{video_id}' não está disponível."}
            except Exception as e:
                logging.error(f"Ocorreu um erro inesperado ao buscar a transcrição: {e}", exc_info=False)
                return {"success": False, "error": f"Ocorreu um erro inesperado ao buscar a transcrição. Detalhe: {e}"}

        tab1, tab2 = st.tabs(["Digitar Texto", "Narrar Vídeo do YouTube"])
        with tab1:
            st.subheader("Opção 1: Digite ou cole seu texto")
            voices = {"Português (Brasil)": {"Feminina": "pt-BR-Standard-A", "Masculina": "pt-BR-Wavenet-B"},
                      "Inglês (EUA)": {"Feminina": "en-US-Standard-C", "Masculina": "en-US-Wavenet-D"},
                      "Espanhol (Espanha)": {"Feminina": "es-ES-Standard-A", "Masculina": "es-ES-Wavenet-B"}}
            text_input_manual = st.text_area("Texto para narrar:", height=200, placeholder="Digite ou cole seu texto aqui...", key="text_manual")
            col1, col2 = st.columns(2)
            with col1:
                language = st.selectbox("Idioma:", list(voices.keys()), key="lang_manual")
            with col2:
                gender = st.selectbox("Gênero da Voz:", list(voices[language].keys()), key="gender_manual")
            voice_name_selected = voices[language][gender]
            speaking_rate = st.slider("Velocidade da Fala:", min_value=0.5, max_value=2.0, value=1.0, step=0.25, key="rate_manual")
            if st.button("🎧 Gerar Áudio do Texto", use_container_width=True, key="btn_manual"):
                if not text_input_manual.strip():
                    st.error("Por favor, insira um texto para gerar o áudio.")
                else:
                    with st.spinner("A IA está gerando sua narração..."):
                        lang_code = voice_name_selected.split('-')[0] + '-' + voice_name_selected.split('-')[1]
                        audio_result = ni_sintetizar_fala(text_input_manual, lang_code, voice_name_selected, speaking_rate)
                    if audio_result["success"]:
                        st.success("Áudio gerado com sucesso!")
                        st.audio(audio_result["data"], format="audio/mp3")
                        st.download_button(label="📥 Baixar Áudio (MP3)", data=audio_result["data"], file_name="narracao_ia.mp3", mime="audio/mp3")
                    else:
                        st.error(f"Falha ao gerar o áudio: {audio_result['error']}")

        with tab2:
            st.subheader("Opção 2: Insira uma URL do YouTube")
            youtube_url = st.text_input("URL do vídeo:", placeholder="Ex: https://www.youtube.com/watch?v=dQw4w9WgXcQ", key="yt_url")
            if st.button("Buscar, Traduzir e Narrar", use_container_width=True):
                st.session_state.pop('translated_transcript_narrador', None)
                st.session_state.pop('audio_bytes_youtube', None)
                st.session_state.pop('narrador_error', None)
                if youtube_url.strip():
                    transcript_result = ni_get_and_translate_transcript(youtube_url, GEMINI_API_KEY)
                    if transcript_result["success"]:
                        st.success("Transcrição traduzida com sucesso!")
                        translated_text = transcript_result["data"]
                        st.session_state['translated_transcript_narrador'] = translated_text
                        st.info("Gerando a narração do vídeo...")
                        narration_voice = "pt-BR-Wavenet-B"
                        narration_rate = 1.0
                        audio_result = ni_sintetizar_fala(translated_text, "pt-BR", narration_voice, narration_rate)
                        if audio_result["success"]:
                            st.success("Narração do vídeo gerada!")
                            st.session_state['audio_bytes_youtube'] = audio_result["data"]
                        else:
                            st.session_state['narrador_error'] = f"A tradução funcionou, mas a geração de áudio falhou: {audio_result['error']}"
                    else:
                        st.session_state['narrador_error'] = f"Falha na etapa de transcrição/tradução: {transcript_result['error']}"
                else:
                    st.warning("Por favor, insira uma URL do YouTube.")
            if 'narrador_error' in st.session_state:
                st.error(st.session_state.narrador_error)
            if 'translated_transcript_narrador' in st.session_state:
                with st.expander("Ver Texto Traduzido", expanded=True):
                    st.text_area("Conteúdo do vídeo em Português:", value=st.session_state.translated_transcript_narrador, height=300, key="edited_transcript_area")
            if 'audio_bytes_youtube' in st.session_state:
                st.subheader("Áudio da Narração")
                st.audio(st.session_state.audio_bytes_youtube, format="audio/mp3")
                st.download_button(label="📥 Baixar Narração do Vídeo (MP3)", data=st.session_state.audio_bytes_youtube, file_name="narracao_video.mp3", mime="audio/mp3")

    def page_estudio_de_mistura_visual():
        st.header("🎨 Estúdio de Mistura Visual")
        st.markdown("Transforme, adapte ou misture imagens para criar memes, caricaturas ou arte digital única.")
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        GEMINI_API_KEY = get_api_key('gemini_key')
        if not GCP_PROJECT_ID or not GCP_LOCATION or not GEMINI_API_KEY:
            st.warning("As chaves do Google Cloud e Gemini são necessárias.")
            return

        def emv_generate_image_from_image(image_bytes, prompt, style, aspect_ratio="1:1"):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                with st.status("Analisando imagem...", expanded=True) as status:
                    multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                    image_part = Part.from_data(data=image_bytes, mime_type='image/png')
                    prompt_analise = f"""
                    Você é um engenheiro de prompts especialista. Analise a imagem e crie um prompt em inglês, detalhado, para um modelo de IA de imagem.
                    - Descreva a base: {prompt}
                    - Aplique o estilo: {style}
                    """
                    response_analysis = multimodal_model.generate_content([prompt_analise, image_part])
                    optimized_prompt = response_analysis.text
                    status.update(label="Análise concluída!", state="complete")
                with st.expander("Ver prompt técnico"): st.code(optimized_prompt, language='text')
                with st.spinner("Gerando nova imagem..."):
                    model_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                    response_image = model_imagem.generate_images(prompt=optimized_prompt, number_of_images=1, aspect_ratio=aspect_ratio)
                    if response_image and response_image.images and response_image.images[0]._image_bytes:
                        return {"success": True, "data": response_image.images[0]._image_bytes}
                    return {"success": False, "error": "IA se recusou a gerar a imagem."}
            except Exception as e:
                return {"success": False, "error": f"Erro na geração da imagem: {e}"}
        
        def create_image_generation_ui(tab_name):
            defaults = {
                "fusao": {"prompt": "Transforme em uma cena épica.", "style": "arte conceitual", "button": "🚀 Iniciar Fusão"},
                "caricatura": {"prompt": "Transforme em uma caricatura engraçada.", "style": "desenho animado", "button": "🖌️ Estilizar"},
                "meme": {"prompt": "Crie um meme com esta imagem.", "style": "estilo de meme da internet", "button": "😂 Gerar Meme"}
            }
            st.subheader(f"Opção: {tab_name.replace('_', ' ').title()}")
            uploaded_image = st.file_uploader("1. Envie sua imagem:", type=["png", "jpg"], key=f"{tab_name}_upload")
            prompt = st.text_area("2. Descreva a transformação:", defaults[tab_name]["prompt"], key=f"{tab_name}_prompt")
            style = st.text_input("3. Estilo artístico:", defaults[tab_name]["style"], key=f"{tab_name}_style")
            if st.button(defaults[tab_name]["button"], use_container_width=True, key=f"{tab_name}_btn"):
                if uploaded_image and prompt:
                    with st.spinner("Processando..."):
                        result = emv_generate_image_from_image(uploaded_image.getvalue(), prompt, style)
                    if result["success"]:
                        st.success("Imagem criada!")
                        st.image(result["data"], use_container_width=True)
                        st.download_button("📥 Baixar Imagem", result["data"], f"{tab_name}.png", "image/png")
                    else:
                        st.error(f"Falha: {result['error']}")
                else:
                    st.error("Por favor, envie uma imagem e descreva a transformação.")
        
        tab1, tab2, tab3 = st.tabs(["🖼️ Fusão", "🎨 Caricatura", "😂 Meme"])
        with tab1: create_image_generation_ui("fusao")
        with tab2: create_image_generation_ui("caricatura")
        with tab3: create_image_generation_ui("meme")

    def page_meus_arquivos():
        st.header("🗂️ Meus Arquivos")
        st.markdown("Aqui estão os arquivos que você gerou e salvou na plataforma.")

        user_dir = Path(__file__).parent / "user_files" / username
        if not user_dir.exists() or not any(user_dir.iterdir()):
            st.info("Você ainda não salvou nenhum arquivo.")
            return

        files = os.listdir(user_dir)
        if not files:
                st.info("Você ainda não salvou nenhum arquivo.")
                return

        for file_name in files:
            file_path = user_dir / file_name
            with open(file_path, "rb") as f:
                st.download_button(
                    label=f"📥 Baixar {file_name}",
                    data=f,
                    file_name=file_name,
                    key=f"download_{file_name}"
                )
            st.markdown("---")
            
    def page_perfil_configuracoes():
        st.header("👤 Perfil e Configurações")
        st.subheader("Gerencie suas chaves de API")
        st.info("Suas chaves são armazenadas de forma segura e usadas para alimentar as ferramentas da plataforma.")

        with st.form("api_keys_form"):
            st.write("**Chaves do Google / Gemini**")
            gemini_key = st.text_input("Gemini API Key", value=get_api_key('gemini_key') or "", type="password")
            gcp_project_id = st.text_input("Google Cloud Project ID", value=get_api_key('gcp_project_id') or "")
            gcp_location = st.text_input("Google Cloud Location", value=get_api_key('gcp_location') or "")

            st.write("**Chaves do Google Custom Search**")
            gsearch_key = st.text_input("Google Custom Search API Key", value=get_api_key('gsearch_key') or "", type="password")
            gsearch_cx = st.text_input("Google Custom Search CX ID", value=get_api_key('gsearch_cx') or "")

            submitted = st.form_submit_button("Salvar Configurações")
            if submitted:
                user_credentials = config['credentials']['usernames'][username]
                user_credentials['api_keys']['gemini_key'] = gemini_key
                user_credentials['api_keys']['gcp_project_id'] = gcp_project_id
                user_credentials['api_keys']['gcp_location'] = gcp_location
                user_credentials['api_keys']['gsearch_key'] = gsearch_key
                user_credentials['api_keys']['gsearch_cx'] = gsearch_cx

                with open(config_file, 'w') as file:
                    yaml.dump(config, file, default_flow_style=False)
                st.success("Suas configurações foram salvas com sucesso!")
                st.rerun()

    # --- ROTEADOR DE PÁGINAS ---
    if page == "Página Inicial":
        page_inicial()
    elif page == "Gerador de Exercícios":
        page_gerador_exercicios()
    elif page == "Otimizador de Prompt":
        page_otimizador_prompt()
    elif page == "Análise Visual de Imagens":
        page_analise_visual()
    elif page == "Criador de Aplicativos":
        page_criador_aplicativos()
    elif page == "Fábrica de Spritesheets 2D":
        page_fabrica_spritesheets()
    elif page == "Análise de Logs":
        page_analise_logs()
    elif page == "Gerador POP Retrogaming": # <-- LÓGICA ADICIONADA AO ROTEADOR
        page_criador_pop_retrogaming()
    elif page == "Espelho da Mente":
        page_espelho_da_mente()
    elif page == "Pesquisa Avançada (Dorks)":
        page_pesquisa_avancada()
    elif page == "Narrador IA":
        page_narrador_ia()
    elif page == "Estúdio de Mistura Visual":
        page_estudio_de_mistura_visual()        
    elif page == "Meus Arquivos":
        page_meus_arquivos()
    elif page == "Perfil e Configurações":
        page_perfil_configuracoes()