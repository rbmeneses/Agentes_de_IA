# coding: utf-8
import streamlit as st
import os
import requests
import logging
import base64
import json
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import tempfile
import vertexai
from vertexai.vision_models import ImageGenerationModel
from vertexai.generative_models import GenerativeModel, Part
from PIL import Image as PILImage
from googleapiclient.discovery import build
from google.api_core import exceptions as google_exceptions
import yaml
from yaml.loader import SafeLoader
import streamlit_authenticator as stauth
from pathlib import Path

# --- CONFIGURAÇÃO GERAL E DA PÁGINA ---
st.set_page_config(layout="wide", page_title="Minha Plataforma de IA ")

# --- 1. DESIGN E ESTILO (UX/UI) ---
# O CSS foi mantido para preservar a identidade visual da aplicação.
st.markdown("""
<style>
    /* Gradiente na barra lateral */
    [data-testid="stSidebar"] {
        background: linear-gradient(135deg, #3B82F6, #8B5CF6);
        color: white;
    }
    [data-testid="stSidebar"] .st-emotion-cache-1629p8f a, [data-testid="stSidebar"] .st-emotion-cache-10trblm {
        color: white;
    }
    /* Estilo dos botões */
    .stButton>button {
        background-color: #8B5CF6;
        color: white;
        border: none;
        border-radius: 8px;
        padding: 10px 20px;
        transition: background-color 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #6D28D9;
    }
    /* Tipografia e Layout */
    body {
        font-family: 'Lato', 'Roboto', sans-serif;
    }
    h1, h2, h3 {
        font-weight: 300; /* Fontes mais leves e elegantes */
    }
    .stTabs, .stRadio {
        margin-top: 2em; /* Espaçamento generoso */
    }
</style>
""", unsafe_allow_html=True)


# --- 2. SISTEMA DE AUTENTICAÇÃO E GESTÃO DE USUÁRIO ---
# A estrutura de autenticação foi mantida, pois é robusta.
config_file = Path(__file__).parent / "config.yaml"
if not config_file.exists():
    # Cria um arquivo de configuração padrão se não existir
    default_config = {
        "credentials": {
            "usernames": {
                "admin": {
                    "email": "admin@example.com",
                    "name": "Administrador",
                    # Senha '12345' - Use uma ferramenta para gerar um hash bcrypt seguro para produção
                    "password": "$2b$12$Yg.i2h.fA94ccbCo3x.iU.W3Yv8M2d2yVpr0dFzgehJe.eAIqfC6C",
                    "api_keys": {
                        "gemini_key": "", "gcp_project_id": "", "gcp_location": "us-central1",
                        "gsearch_key": "", "gsearch_cx": ""
                    }
                }
            }
        },
        "cookie": {"expiry_days": 30, "key": "some_signature_key", "name": "ia_plataforma_cookie"},
        "preauthorized": {"emails": ["admin@example.com"]}
    }
    with open(config_file, 'w') as file:
        yaml.dump(default_config, file, default_flow_style=False)

with open(config_file) as file:
    config = yaml.load(file, Loader=SafeLoader)

authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)

authenticator.login()

name = st.session_state.get("name")
authentication_status = st.session_state.get("authentication_status")
username = st.session_state.get("username")

# --- LÓGICA PRINCIPAL DA APLICAÇÃO ---

if not authentication_status:
    st.warning("Por favor, faça login para acessar a plataforma.")
    if authentication_status is False:
        st.error('Usuário ou senha incorretos.')
    elif authentication_status is None:
        st.info('Bem-vindo! Por favor, insira seu usuário e senha.')

elif authentication_status:
    # --- FUNÇÕES UTILITÁRIAS GLOBAIS E CONSTANTES ---
    # COMENTÁRIO DE REFATORAÇÃO: Centralização de constantes para fácil manutenção.
    GEMINI_API_BASE_URL = "https://generativelanguage.googleapis.com/v1beta"
    GEMINI_FLASH_MODEL = "gemini-2.5-flash"
    GEMINI_PRO_MODEL = "gemini-2.5-pro" # Updated to a robust multimodal model for log analysis
    VERTEX_IMAGE_MODEL = "imagen-3.0-fast-generate-001"

    def get_api_key(key_name):
        """Busca a chave de API do perfil do usuário logado de forma segura."""
        try:
            return config['credentials']['usernames'][username]['api_keys'][key_name]
        except (KeyError, TypeError):
            st.error(f"Chave de API '{key_name}' não encontrada. Configure-a na página 'Perfil e Configurações'.")
            return None

    def save_file_to_user_storage(file_stream, filename):
        """Salva um arquivo na área de armazenamento persistente do usuário."""
        user_dir = Path(__file__).parent / "user_files" / username
        user_dir.mkdir(parents=True, exist_ok=True)
        file_stream.seek(0)
        with open(user_dir / filename, "wb") as f:
            f.write(file_stream.getbuffer())
        st.success(f"Arquivo '{filename}' salvo com sucesso em 'Meus Arquivos'!")

    # *** CORREÇÃO APLICADA AQUI ***
    # A função foi movida para o escopo global para ser reutilizável.
    def get_gemini_response(prompt_text, model_name, temperature, api_key):
        """Envia um prompt para a API Gemini e retorna a resposta em texto."""
        if not api_key:
            st.error("API Key for Gemini is not configured.")
            return None
        api_endpoint = f"{GEMINI_API_BASE_URL}/models/{model_name}:generateContent?key={api_key}"
        payload = {
            "contents": [{"parts": [{"text": prompt_text}]}],
            "generationConfig": {"temperature": temperature},
            "safetySettings": [
                {
                    "category": "HARM_CATEGORY_HARASSMENT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_HATE_SPEECH",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                }
            ],
        }
        try:
            response = requests.post(api_endpoint, json=payload, headers={"Content-Type": "application/json"})
            response.raise_for_status()
            result = response.json()
            if 'candidates' in result and result['candidates'][0]['content']['parts'][0]['text']:
                return result['candidates'][0]['content']['parts'][0]['text']
            else:
                st.warning("A resposta da IA não continha o texto esperado.")
                return f"Resposta inesperada da API: {result}"
        except requests.exceptions.RequestException as e:
            st.error(f"Erro de requisição com a API Gemini: {e}")
            return f"Erro ao comunicar com a API Gemini: {e}"
        except Exception as e:
            st.error(f"Erro ao processar o prompt: {e}")
            return f"Erro ao processar o prompt: {e}"

    # --- BARRA LATERAL DE NAVEGAÇÃO ---
    with st.sidebar:
        st.title(f"Bem-vindo, {name}")
        st.markdown("---")

        page = st.radio("Selecione uma Ferramenta:",
                        ("Página Inicial", "Gerador de Exercícios", "Otimizador de Prompt",
                         "Análise Visual de Imagens", "Criador de Aplicativos",
                         "Fábrica de Spritesheets 2D", "Análise de Logs",
                         "Espelho da Mente", "Buscador de Vagas", # Consolidated job search
                         "Meus Arquivos", "Perfil e Configurações"))

        st.markdown("---")
        authenticator.logout("Logout", "main")

    # --- DEFINIÇÃO DAS PÁGINAS ---
    def page_inicial():
        st.title("🚀 Minhas Ferramentas de IA")
        st.markdown("### Bem-vindo à sua central de ferramentas de Inteligência Artificial.")
        st.success("**Novidade:** Confira os novos módulos 'Análise Visual de Imagens', 'Criador de Aplicativos' e 'Análise de Logs'!")
        st.info("Navegue pelas ferramentas usando o menu à esquerda. Configure suas chaves de API na página 'Perfil e Configurações' para habilitar todas as funcionalidades.")

    def page_gerador_exercicios():
        st.header("🧩 Gerador de Exercícios para Estudo Adaptados")
        st.markdown("Crie exercícios personalizados. A geração de imagens requer a configuração do seu Projeto Google Cloud.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location') # Ensure GCP_LOCATION is also retrieved
        if not GEMINI_API_KEY: return
        if not GCP_PROJECT_ID: return # Added explicit check for project ID
        if not GCP_LOCATION: return # Added explicit check for location

        def ge_extract_image_prompts(content, desired_image_style=""):
            parts = content.split("[IMAGEM]")
            image_prompts = []
            for i in range(len(parts) - 1):
                before_context = parts[i][-150:]
                after_context = parts[i + 1][:150]
                image_description = f"{before_context.strip()} {after_context.strip()}"
                if desired_image_style:
                    prompt = (f"Uma imagem educacional no estilo '{desired_image_style}' para ilustrar o seguinte conceito: {image_description}. "
                              "A imagem deve ser clara, didática e focada no objeto principal.")
                else:
                    prompt = f"Uma imagem educacional detalhada para o seguinte contexto: {image_description}"
                image_prompts.append(prompt)
            return image_prompts

        def ge_gerar_imagem_com_vertexai(prompt):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION) # Initialize Vertex AI here
                modelo_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo_imagem.generate_images(prompt=prompt, number_of_images=1)
                if not resposta or not resposta[0]._image_bytes:
                    st.warning(f"A IA se recusou a gerar a imagem para o prompt: '{prompt}'.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Ocorreu um erro ao gerar a imagem via Vertex AI: {e}")
                return None

        def generate_exercises(data, api_key):
            try:
                gemini_text_model = GEMINI_FLASH_MODEL
                question_prompt = (
                    f"Por favor, gere {data['questionCount']} exercícios do tipo {data['questionType']} para o nível escolar {data['gradeLevel']}, "
                    f"dificuldade {data['difficulty']} sobre o tema '{data['theme']}'.\n"
                    f"Adapte o conteúdo para {data['specialNeed']}.\n"
                )
                if data['include_images']:
                    question_prompt += ("\n**Instrução Crítica para Imagens:** Você DEVE inserir o marcador `[IMAGEM]` no texto em locais relevantes para ilustrar conceitos-chave. "
                                      "É obrigatório que o marcador `[IMAGEM]` apareça no texto gerado. Exemplo: '...a mitocôndria, que é a usina de energia da célula. [IMAGEM]'\n")

                gemini_payload = {"contents": [{"parts": [{"text": question_prompt}]}]}
                url = f"{GEMINI_API_BASE_URL}/models/{gemini_text_model}:generateContent?key={api_key}"
                headers = {"Content-Type": "application/json"}
                response = requests.post(url, json=gemini_payload, headers=headers)
                response.raise_for_status()
                gemini_result = response.json()
                content = "".join(part.get('text', '') for part in gemini_result['candidates'][0]['content'].get('parts', []))

                images_data = []
                if data['include_images']:
                    image_prompts = ge_extract_image_prompts(content, data['imageStyle'])
                    if image_prompts:
                        with st.spinner(f"Gerando {len(image_prompts)} imagens..."):
                            for prompt_text in image_prompts:
                                img_bytes = ge_gerar_imagem_com_vertexai(prompt_text)
                                if img_bytes: images_data.append(img_bytes)

                file_stream = BytesIO()
                if data['outputFormat'] == 'docx':
                    doc = Document()
                    doc.add_heading('Exercícios Gerados', 0)
                    text_parts = content.split('[IMAGEM]')
                    for i, part_text in enumerate(text_parts):
                        doc.add_paragraph(part_text.strip())
                        if i < len(images_data):
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                                temp_img.write(images_data[i])
                                temp_img_path = temp_img.name
                            doc.add_picture(temp_img_path, width=doc.sections[0].page_width * 0.5)
                            os.remove(temp_img_path)
                    doc.save(file_stream)
                    file_stream.seek(0)
                    return file_stream, 'exercicios.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                else: # Handles txt and pdf
                    if data['outputFormat'] == 'pdf':
                        c = canvas.Canvas(file_stream, pagesize=A4)
                        textobject = c.beginText()
                        textobject.setTextOrigin(10, 750)
                        textobject.setFont("Helvetica", 12)
                        lines = content.split('\n')
                        for line in lines:
                            textobject.textLine(line)
                        c.drawText(textobject)
                        c.showPage()
                        c.save()
                        file_stream.seek(0)
                        return file_stream, f'exercicios.pdf', 'application/pdf'
                    else: # txt
                        file_stream.write(content.encode('utf-8'))
                        file_stream.seek(0)
                        return file_stream, f'exercicios.txt', 'text/plain'
            except Exception as e:
                raise e

        with st.form(key="exercise_form"):
            theme_input = st.text_area("📝 Tema ou Instruções:", "Fotossíntese para o ensino fundamental", height=100)
            col1, col2, col3 = st.columns(3)
            with col1:
                question_count = st.slider("🔢 Número de Exercícios", 1, 30, 5)
                grade_level = st.selectbox("🎓 Nível Escolar", ["Infantil", "Fundamental I", "Fundamental II", "Médio", "Superior"], index=2)
            with col2:
                difficulty = st.slider("📊 Dificuldade (0-100)", 0, 100, 40)
                question_type = st.selectbox("✍️ Tipo de Questão", ["Discursiva", "Múltipla Escolha"])
            with col3:
                include_images = st.checkbox("🖼️ Incluir imagens com Vertex AI", value=False)
                output_format = st.selectbox("💾 Formato de Saída", ["docx", "txt", "pdf"], help="Escolha DOCX para incluir imagens; PDF e TXT não incluirão imagens.")
                image_style = st.text_input("🎨 Estilo da Imagem", "desenho vetorial simples", disabled=not include_images)

            special_need = st.selectbox("♿ Necessidade Específica (Opcional)", ["Nenhuma necessidade específica", "Síndrome de Down", "TEA", "Deficiência Intelectual", "TDAH"])
            submit_button = st.form_submit_button("🚀 Gerar Exercícios")

        if submit_button:
            if not theme_input.strip():
                st.error("Por favor, insira um tema.")
            else:
                with st.spinner("Gerando exercícios..."):
                    try:
                        if include_images:
                            if not GCP_PROJECT_ID or not GCP_LOCATION:
                                st.error("Project ID ou Location do Google Cloud não configurado no perfil para gerar imagens.")
                                return
                            # vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION) # Moved to ge_gerar_imagem_com_vertexai

                        payload = {'theme': theme_input, 'questionCount': question_count,
                                   'gradeLevel': grade_level, 'difficulty': difficulty, 'questionType': question_type,
                                   'specialNeed': special_need, 'include_images': include_images,
                                   'imageStyle': image_style, 'outputFormat': output_format}

                        file_stream, filename, mime = generate_exercises(payload, GEMINI_API_KEY)

                        st.success("✅ Exercícios gerados com sucesso!")
                        col1, col2 = st.columns(2)
                        with col1:
                           st.download_button(label="📥 Baixar Arquivo", data=file_stream, file_name=filename, mime=mime, use_container_width=True)
                        with col2:
                            if st.button("💾 Salvar em Meus Arquivos", use_container_width=True):
                                save_file_to_user_storage(file_stream, filename)

                    except Exception as e:
                        st.error(f"❌ Erro ao gerar os exercícios: {e}")

    def page_otimizador_prompt():
        st.header("✨ Gerador de Prompts Otimizados com Gemini")
        st.markdown("Preencha os campos abaixo para usar a IA para gerar um prompt otimizado para outra IA.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        if not GEMINI_API_KEY: return

        media_type = st.selectbox("Tipo de Mídia:", ("Texto", "Imagem", "Vídeo"), key="op_media")
        request_type = st.selectbox("Tipo de Requisição:", ("Geração de Conteúdo", "Resumo", "Tradução", "Análise de Sentimento", "Ideação", "Geração de Código", "Debug", "Refatoração", "Documentação", "Outro"), key="op_req")
        specific_details = st.text_area("Detalhes Específicos (tom, formato, público):", placeholder="Ex: Tom formal, formato de lista, público: desenvolvedores.", key="op_details")
        content = st.text_area("Requisitos do Usuário (o que você deseja que a IA faça):", placeholder="Ex: Escreva um artigo sobre os benefícios da IA para pequenas empresas.", height=150, key="op_content")
        example = st.text_area("Exemplo (opcional):", placeholder="Ex: 'Introdução: A Inteligência Artificial está transformando...'", key="op_example")
        if media_type in ("Imagem", "Vídeo"):
            media_details = st.text_area("Detalhes da Imagem/Vídeo:", placeholder="Ex: Estilo realista, plano próximo, 4K.", key="op_media_details")
        else:
            media_details = ""

        gemini_model = st.selectbox("Modelo Gemini:", (GEMINI_FLASH_MODEL, GEMINI_PRO_MODEL), key="op_model")
        temperature = st.slider("Temperatura:", 0.0, 1.0, 0.7, 0.05, key="op_temp")

        if st.button("Gerar Prompt Otimizado", use_container_width=True):
            if not content.strip():
                st.error("O campo 'Requisitos do Usuário' é obrigatório.")
            else:
                with st.spinner("Otimizando seu prompt..."):
                    prompt_for_gemini = (f"Gere um prompt otimizado para {media_type} com base nas informações: "
                                         f"- Requisição: {request_type} - Detalhes: {specific_details} - Requisitos: {content} "
                                         f"- Exemplo: {example or 'N/A'} - Detalhes de Mídia: {media_details or 'N/A'}")
                    # *** CORREÇÃO APLICADA AQUI ***
                    # A chamada foi atualizada para a nova função global.
                    enhanced_prompt = get_gemini_response(prompt_for_gemini, gemini_model, temperature, GEMINI_API_KEY)
                    if enhanced_prompt:
                        st.subheader("Prompt Otimizado Gerado:")
                        st.code(enhanced_prompt, language='text')

    def page_analise_visual():
        st.header("👁️‍🗨️ Análise Visual para Geração de Conteúdo")
        st.markdown("Envie uma imagem (diagrama, UI, esquema, etc.) e a IA irá gerar código e um prompt baseado nela.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GEMINI_API_KEY or not GCP_PROJECT_ID or not GCP_LOCATION: return

        def av_analisar_imagem(image_bytes, prompt_usuario):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL) # Use Pro model for multimodal capabilities
                image_part = Part.from_data(image_bytes, mime_type="image/png")

                instrucao_ia = f"""
                Analise a imagem fornecida. Com base em seu conteúdo visual, gere duas saídas distintas e claramente separadas:

                1.  **Geração de Código:** Crie um trecho de código relevante (Python, HTML, etc.) ou pseudocódigo que represente a lógica, estrutura ou os elementos visuais da imagem. Se for um diagrama, gere a classe correspondente. Se for um layout de UI, gere o HTML/CSS básico. Adicione um comentário no início do bloco de código indicando a linguagem.

                2.  **Geração de Prompt:** Crie um prompt de texto conciso e otimizado que descreva a imagem. Este prompt deve ser formatado para ser usado diretamente na ferramenta 'Otimizador de Prompt', capturando a essência da imagem para uma futura geração de conteúdo.

                Se o usuário fornecer uma instrução específica, leve-a em consideração: '{prompt_usuario}'
                """

                response = multimodal_model.generate_content([instrucao_ia, image_part])
                return response.text

            except google_exceptions.GoogleAPIError as e:
                st.error(f"Google API Error durante análise de imagem: {e.message}")
                return None
            except Exception as e:
                st.error(f"Erro ao analisar a imagem com o modelo multimodal: {e}")
                return None

        uploaded_file = st.file_uploader("Selecione uma imagem", type=["png", "jpg", "jpeg"])
        user_prompt = st.text_input("Instrução Específica (Opcional):", placeholder="Ex: 'Foco no formulário de login', 'Gere o código em Python'")

        if uploaded_file is not None:
            st.image(uploaded_file, caption="Imagem Carregada.", use_column_width=True)

            if st.button("Analisar Imagem e Gerar Conteúdo", use_container_width=True):
                with st.spinner("A IA está analisando a imagem..."):
                    try:
                        image_data = uploaded_file.getvalue()

                        resultado_analise = av_analisar_imagem(image_data, user_prompt)

                        if resultado_analise:
                            st.success("Análise concluída com sucesso!")
                            st.markdown("---")

                            partes = resultado_analise.split("Geração de Prompt:")
                            if len(partes) == 2:
                                st.subheader("Código Gerado a partir da Imagem:")
                                st.code(partes[0].replace("Geração de Código:", "").strip(), language='python') # Default to python, user can change if needed
                                st.subheader("Prompt Otimizado para Descrever a Imagem:")
                                st.code(partes[1].strip(), language='text')
                            else:
                                st.subheader("Saída da IA (formato inesperado):")
                                st.code(resultado_analise, language='text')

                    except Exception as e:
                        st.error(f"Ocorreu um erro no processo: {e}")

    def page_criador_aplicativos():
        st.header("🏗️ Criador de Aplicativos (Scaffolding)")
        st.markdown("Gere a estrutura de arquivos e o código base para diferentes tipos de aplicações. Agora com prompts otimizados!")

        def ca_gerar_scaffold(project_name, app_type, extra_functions=None):
            templates = {
                "Streamlit App Simples": {
                    "app.py": f"""
import streamlit as st

st.set_page_config(page_title="{project_name}")

def main():
    st.title("Bem-vindo ao {project_name}!")
    st.write("Este é um aplicativo Streamlit simples gerado pela Plataforma IA Evoluída.")

    name = st.text_input("Qual é o seu nome?")
    if name:
        st.write(f"Olá, {{name}}!")
{extra_functions or ""}
if __name__ == "__main__":
    main()
                    """,
                    "requirements.txt": "streamlit"
                },
                "API Flask Básica": {
                    "app.py": f"""
from flask import Flask, jsonify

app = Flask(__name__)

@app.route('/')
def index():
    return jsonify({{"message": "Bem-vindo à API '{project_name}'!"}})

@app.route('/api/data')
def get_data():
    return jsonify({{"id": 1, "name": "Dado de Exemplo"}})
{extra_functions or ""}
if __name__ == '__main__':
    app.run(debug=True)
                    """,
                    "requirements.txt": "Flask"
                },
                "Script de Automação": {
                     "main.py": f"""
import os

def main():
    print("Iniciando o script de automação: {project_name}")
    # Exemplo: Listar arquivos no diretório atual
    files = os.listdir('.')
    print("Arquivos encontrados:")
    for f in files:
        print(f"- {{f}}")
    print("Script concluído.")
{extra_functions or ""}
if __name__ == "__main__":
    main()
                    """,
                    "README.md": f"# {project_name}\\n\\nEste é um script de automação gerado pela Plataforma IA Evoluída."
                }
            }
            return templates.get(app_type, {})

        st.subheader("1. Defina os Detalhes do Projeto")
        project_name_input = st.text_input("Nome do Projeto:", "MeuNovoApp")
        app_type_select = st.selectbox(
            "Selecione o Tipo de Aplicativo:",
            ("Streamlit App Simples", "API Flask Básica", "Script de Automação")
        )

        st.subheader("2. Adicione Funções Extras com IA")
        extra_func_description = st.text_area(
            "Descreva funções extras para adicionar ao aplicativo (ex: exportar dados, autenticação, gráficos):",
            placeholder="Ex: Adicionar função para exportar dados em CSV."
        )

        GEMINI_API_KEY = get_api_key('gemini_key')
        gemini_model = GEMINI_FLASH_MODEL
        temperature = 0.7

        extra_functions_code = ""
        if extra_func_description.strip():
            if not GEMINI_API_KEY:
                st.warning("Chave Gemini API não configurada. Funções extras via IA não serão geradas.")
            else:
                with st.spinner("Otimizando prompt e gerando funções extras..."):
                    prompt_for_gemini = (
                        f"Gere funções Python para {app_type_select} conforme a descrição: {extra_func_description}. "
                        "O código deve ser pronto para uso e fácil de integrar ao template padrão."
                    )
                    # *** CORREÇÃO APLICADA AQUI ***
                    # A chamada foi atualizada para a nova função global, resolvendo o NameError.
                    enhanced_code = get_gemini_response(prompt_for_gemini, gemini_model, temperature, GEMINI_API_KEY)
                    if enhanced_code:
                        extra_functions_code = "\n" + enhanced_code

        if st.button("Gerar Estrutura do Aplicativo", use_container_width=True):
            if not project_name_input.strip():
                st.error("Por favor, forneça um nome para o projeto.")
            else:
                with st.spinner("Gerando arquivos do projeto..."):
                    scaffold_files = ca_gerar_scaffold(project_name_input, app_type_select, extra_functions=extra_functions_code)
                    if scaffold_files:
                        st.success(f"Estrutura para '{project_name_input}' gerada!")
                        st.balloons()
                        for filename, content in scaffold_files.items():
                            with st.expander(f"📄 Arquivo: {filename}"):
                                st.code(content, language='python' if filename.endswith('.py') else 'text')
                    else:
                        st.error("Tipo de aplicativo não encontrado.")

    def page_fabrica_spritesheets():
        st.header("👹 Fábrica de Spritesheets 2D - Geração por Template Avançado")
        st.markdown("Utiliza um template de prompt de alta qualidade para gerar a spritesheet completa de uma só vez.")

        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GCP_PROJECT_ID or not GCP_LOCATION: return

        def fs_gerar_prompt_template_avancado(data, lista_frames, estilo_arte):
            style_text = ("2D pixel art spritesheet" if estilo_arte == "Pixel Art" else "2D cartoon spritesheet")
            character_text = (f"a powerful '{data['funcao']}' character, wearing a {data['cor']}. "
                              f"Specific details: {data['detalhes']}. Wielding a {data['arma']}.")
            animation_text = ", ".join(lista_frames)
            prompt_final = (f"{style_text}, of {character_text}. "
                            f"The spritesheet must include a complete set of animation frames: {animation_text}. "
                            "Present the sprites in a clear, organized grid layout with a transparent background. "
                            "The character design must be strictly consistent across all poses. "
                            "Use a side-view perspective, optimized for smooth game animation. High resolution image.")
            return prompt_final

        def fs_gerar_imagem_google(prompt):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION) # Initialize Vertex AI here
                st.info(f"Enviando o prompt para a IA...")
                modelo = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo.generate_images(prompt=prompt, number_of_images=1, aspect_ratio="16:9")
                if not resposta or not resposta[0]._image_bytes:
                    st.warning("A IA se recusou a gerar a imagem. Tente uma descrição diferente.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Erro ao gerar a imagem via Vertex AI: {e}")
                return None

        st.subheader("1. Defina os Detalhes do Personagem")
        estilo_arte_selecionado = st.selectbox("Estilo de Arte", ["Pixel Art", "Cartoon"], key="fs_style_select")
        col1, col2 = st.columns(2)
        with col1:
            funcao = st.text_input("Função ou Classe", "Mago Arcano", key="fs_funcao")
            arma = st.text_input("Arma/Equipamento", "cajado de madeira com orbe brilhante", key="fs_arma")
        with col2:
            cor = st.text_input("Cores Dominantes", "manto azul escuro e túnica roxa", key="fs_cor")
        detalhes = st.text_area("Detalhes Visuais", "ancião com longa barba branca e expressão sábia", key="fs_detalhes", height=100)

        st.subheader("2. Selecione as Animações")
        frames_disponiveis = ["Idle", "Walk", "Run", "Jump (take-off and landing)", "Basic Melee Attack (staff swing)", "Magic Spell Cast (fireball)", "Hurt", "Block", "Death animation"]
        frames_selecionados = st.multiselect("Selecione os frames para a spritesheet", options=frames_disponiveis, default=["Idle", "Walk", "Run", "Magic Spell Cast (fireball)"])

        if st.button("👹 Gerar Spritesheet", use_container_width=True):
            if not frames_selecionados:
                st.warning("Selecione pelo menos um frame de animação.")
            else:
                try:
                    # vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION) # Moved to fs_gerar_imagem_google
                    with st.spinner("Construindo prompt e gerando a imagem..."):
                        dados_personagem = {"funcao": funcao, "detalhes": detalhes, "arma": arma, "cor": cor}
                        prompt_completo = fs_gerar_prompt_template_avancado(dados_personagem, frames_selecionados, estilo_arte_selecionado)
                        imagem_gerada_bytes = fs_gerar_imagem_google(prompt_completo)

                        if imagem_gerada_bytes:
                            st.success("Spritesheet gerada com sucesso!")
                            st.image(imagem_gerada_bytes, caption="Resultado da geração com o template.", use_container_width=True)
                            st.download_button("📥 Baixar Spritesheet (PNG)", data=imagem_gerada_bytes, file_name=f"{funcao.replace(' ', '_').lower()}_spritesheet.png", mime="image/png", use_container_width=True)
                except Exception as e:
                    st.error(f"Erro ao gerar imagens com Vertex AI: {e}")

    # --- NEW PAGE: ANÁLISE DE LOGS ---
    def page_analise_logs():
        st.header("📊 Análise de Logs e Eventos com IA")
        st.markdown("Cole o conteúdo do seu log ou faça upload de uma imagem (ex: screenshot de erro) para a IA analisar e fornecer insights.")

        GEMINI_API_KEY = get_api_key('gemini_key')
        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')

        if not GEMINI_API_KEY or not GCP_PROJECT_ID or not GCP_LOCATION:
            st.warning("Configure suas chaves de API para Gemini, Google Cloud Project ID e Location na aba 'Perfil e Configurações' para usar esta ferramenta.")
            return

        def al_analisar_imagem_multimodal(image_bytes, prompt_usuario):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                multimodal_model = GenerativeModel(GEMINI_PRO_MODEL)
                image_part = Part.from_data(image_bytes, mime_type="image/png") # Assuming PNG for image logs

                instrucao_ia = f"""
                Analise esta imagem que faz parte de um log ou representa um evento do sistema.
                Extraia qualquer texto visível (OCR), identifique diagramas, gráficos ou estruturas visuais.
                Descreva o conteúdo visual de forma concisa e relevante para a análise de logs.
                Se for um diagrama de fluxo, descreva o fluxo. Se for um erro, transcreva-o.
                Sua análise deve ser útil para complementar a compreensão do log textual.
                Instrução específica do usuário: '{prompt_usuario}'
                """
                response = multimodal_model.generate_content([instrucao_ia, image_part])
                return response.text
            except google_exceptions.GoogleAPIError as e:
                st.error(f"Erro na API do Google ao analisar a imagem: {e.message}")
                return f"Erro na análise visual da imagem: {e.message}"
            except Exception as e:
                st.error(f"Ocorreu um erro inesperado ao analisar a imagem: {e}")
                return f"Erro na análise visual da imagem: {e}"

        with st.form(key="log_analysis_form"):
            st.subheader("1. Forneça o Conteúdo do Log")
            log_content = st.text_area("📜 Cole o conteúdo do log aqui:", height=250,
                                       placeholder="Ex: [2023-10-26 14:30:15] ERROR: Falha na conexão com o banco de dados. Erro: Timeout, host=db.example.com")

            uploaded_image_log = st.file_uploader("🖼️ Ou faça upload de uma imagem de log/evento (opcional):", type=["png", "jpg", "jpeg"])
            image_analysis_instruction = st.text_input("Instrução para análise da imagem (opcional):",
                                                        placeholder="Ex: 'Foque no diagrama de rede', 'Transcreva o erro na tela'")

            st.subheader("2. Detalhes Adicionais (Opcional)")
            col1, col2, col3 = st.columns(3)
            with col1:
                log_type = st.selectbox("Tipo de Log:", ["Geral", "Sistema", "Aplicação", "Segurança", "Rede", "Banco de Dados", "Outro"])
            with col2:
                data_inicio = st.date_input("Data de Início:", value=None, key="log_start_date")
            with col3:
                data_fim = st.date_input("Data de Fim:", value=None, key="log_end_date")

            col4, col5 = st.columns(2)
            with col4:
                gravidade = st.selectbox("Gravidade:", ["Qualquer", "DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"])
            with col5:
                data_source = st.text_input("Fonte de Dados (Sistema/Serviço):", placeholder="Ex: Apache, Nginx, Kubernetes, Servidor DB")

            user_instruction = st.text_area("Instrução específica para a IA (como você quer a análise?):", height=100,
                                            placeholder="Ex: 'Foque em anomalias de desempenho', 'Identifique tentativas de acesso não autorizado'")

            submit_button = st.form_submit_button("🚀 Analisar Log com IA")

        if submit_button:
            if not log_content and not uploaded_image_log:
                st.error("Por favor, cole o conteúdo do log ou faça upload de uma imagem.")
            else:
                with st.spinner("Analisando log e eventos..."):
                    try:
                        image_analysis_result = ""
                        if uploaded_image_log:
                            st.info("Processando imagem de log...")
                            image_bytes = uploaded_image_log.getvalue()
                            image_analysis_result = al_analisar_imagem_multimodal(image_bytes, image_analysis_instruction)
                            if image_analysis_result:
                                st.success("Análise visual concluída.")
                            else:
                                st.warning("Não foi possível obter uma análise visual da imagem.")


                        prompt_parts = []
                        prompt_parts.append("Analise o seguinte log e eventos para identificar padrões, anomalias, problemas potenciais e sugerir melhorias. Forneça um relatório detalhado.")
                        prompt_parts.append(f"\n--- Detalhes do Log ---")
                        prompt_parts.append(f"Tipo de Log: {log_type}")
                        if data_inicio: prompt_parts.append(f"Data de Início: {data_inicio}")
                        if data_fim: prompt_parts.append(f"Data de Fim: {data_fim}")
                        prompt_parts.append(f"Gravidade: {gravidade}")
                        if data_source: prompt_parts.append(f"Fonte de Dados: {data_source}")

                        if log_content:
                            prompt_parts.append(f"\n--- Conteúdo Textual do Log ---\n{log_content}")

                        if image_analysis_result:
                            prompt_parts.append(f"\n--- Análise Visual da Imagem ---\n{image_analysis_result}")

                        if user_instruction:
                            prompt_parts.append(f"\n--- Instruções Específicas do Usuário ---\n{user_instruction}")

                        prompt_parts.append(f"\n--- Estrutura do Relatório Requerida ---")
                        prompt_parts.append("O relatório deve ser estruturado nas seguintes seções: ")
                        prompt_parts.append("1. **Sumário Executivo:** Breve resumo das descobertas mais importantes.")
                        prompt_parts.append("2. **Eventos Chave e Linha do Tempo:** Descrição dos eventos mais críticos ou relevantes, com marcadores de tempo se disponíveis no log.")
                        prompt_parts.append("3. **Padrões Identificados:** Quaisquer repetições, tendências ou comportamentos normais/anormais encontrados.")
                        prompt_parts.append("4. **Anomalias e Problemas:** Detalhes sobre erros, exceções, falhas ou comportamentos inesperados.")
                        prompt_parts.append("5. **Impacto Potencial:** Explicação de como os problemas identificados podem afetar o sistema ou os usuários.")
                        prompt_parts.append("6. **Recomendações e Próximos Passos:** Sugestões acionáveis para resolver problemas, otimizar ou prevenir ocorrências futuras.")
                        prompt_parts.append("Use formatação Markdown para clareza.")

                        full_prompt = "\n".join(prompt_parts)

                        analysis_result = get_gemini_response(full_prompt, GEMINI_PRO_MODEL, temperature=0.2, api_key=GEMINI_API_KEY) # Lower temperature for factual analysis

                        if analysis_result:
                            st.success("✅ Análise de Log Concluída!")
                            st.markdown(analysis_result)
                        else:
                            st.error("❌ Não foi possível gerar a análise do log. Tente novamente ou verifique suas chaves de API.")

                    except Exception as e:
                        st.error(f"❌ Ocorreu um erro inesperado durante a análise: {e}")


    def page_espelho_da_mente():
        st.header("✨ Espelho da Mente Dinâmico")
        st.markdown("Transforme pensamentos e sentimentos complexos em arte simbólica.")

        GCP_PROJECT_ID = get_api_key('gcp_project_id')
        GCP_LOCATION = get_api_key('gcp_location')
        if not GCP_PROJECT_ID or not GCP_LOCATION: return

        def em_criar_prompt_avancado_com_ia(dados_usuario):
            GEMINI_API_KEY = get_api_key('gemini_key') # Retrieve API key here as well for this function
            if not GEMINI_API_KEY:
                return None
            modelo_texto = GenerativeModel(GEMINI_FLASH_MODEL)
            prompt_base = f"""
            Sua tarefa é atuar como um engenheiro de prompt especialista em geração de imagens artísticas.
            Converta a descrição do usuário em um prompt em inglês, detalhado e evocativo para um modelo de IA de imagem.
            O objetivo é criar uma metáfora visual do sentimento do usuário.

            **Diretrizes:**
            1. **Estilo e Atmosfera:** O estilo principal é '{dados_usuario['estilo']}'. Descreva a atmosfera usando luz, sombra e a paleta de cores '{dados_usuario['cores']}' para intensificar a emoção.
            2. **Composição:** Cenário: '{dados_usuario['cenario']}'. Ponto focal: '{dados_usuario['objeto']}', que simboliza a experiência.
            3. **Simbolismo:** Incorpore o sentimento '{dados_usuario['pensamento']}' na cena de forma sutil.
            4. **Detalhes Adicionais:** Inclua: '{dados_usuario['detalhes']}'.
            5. **Restrições:** Evite: '{dados_usuario['evitar']}'.

            Gere o prompt final para a IA de imagem.
            """
            try:
                response = modelo_texto.generate_content([prompt_base])
                return response.text.strip()
            except google_exceptions.GoogleAPIError as e:
                st.error(f"Erro na API do Google ao criar prompt para imagem: {e.message}")
                return None
            except Exception as e:
                st.error(f"Ocorreu um erro inesperado ao criar o prompt: {e}")
                return None

        def em_gerar_imagem_google(prompt):
            try:
                vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION)
                modelo_imagem = ImageGenerationModel.from_pretrained(VERTEX_IMAGE_MODEL)
                resposta = modelo_imagem.generate_images(prompt=prompt, number_of_images=1, aspect_ratio="1:1")
                if not resposta or not resposta[0]._image_bytes:
                    st.warning("A IA recusou-se a gerar a imagem. Tente uma descrição diferente.")
                    return None
                return resposta[0]._image_bytes
            except Exception as e:
                st.error(f"Erro ao gerar a imagem via Vertex AI: {e}")
                return None

        st.subheader("🎨 Construa sua Visão")
        pensamento_usuario = st.text_area("O que você está pensando ou sentindo?", "Nostalgia e esperança por um futuro incerto.", height=100, key="em_thought")
        cenario = st.text_input("Se esse sentimento fosse um lugar, qual seria?", "Uma estação de trem antiga e vazia ao amanhecer.", key="em_place")
        objeto = st.text_input("Qual objeto simboliza melhor este momento?", "Um único broto verde crescendo entre os trilhos.", key="em_object")
        estilo = st.selectbox("Em qual estilo de arte?", ("Surrealismo", "Impressionismo", "Arte Conceitual", "Abstrato", "Aquarela", "Fantasia Sombria", "Cyberpunk"), key="em_style")
        cores = st.text_input("Quais as cores ou paleta?", "Tons de sépia, cinza, com um ponto de luz dourada.", key="em_colors")
        detalhes = st.text_input("Algum detalhe adicional?", "Névoa baixa no chão, raios de sol por uma janela quebrada.", key="em_details_add")
        evitar = st.text_input("O que NÃO incluir na imagem?", "Pessoas, animais, relógios.", key="em_avoid")

        if st.button("✨ Gerar Imagem da Emoção", use_container_width=True):
            if not all([pensamento_usuario, cenario, objeto]):
                st.warning("Por favor, preencha os campos de sentimento, lugar e objeto.")
            else:
                try:
                    # vertexai.init(project=GCP_PROJECT_ID, location=GCP_LOCATION) # Moved to em_gerar_imagem_google
                    with st.spinner("Etapa 1: Interpretando seus sentimentos com IA..."):
                        dados_para_prompt = {"pensamento": pensamento_usuario, "cenario": cenario, "objeto": objeto, "estilo": estilo, "cores": cores, "detalhes": detalhes, "evitar": evitar}
                        prompt_final = em_criar_prompt_avancado_com_ia(dados_para_prompt)

                    if prompt_final:
                        with st.expander("Ver o prompt técnico gerado"):
                            st.write(prompt_final)

                        with st.spinner("Etapa 2: Pintando sua visão..."):
                            imagem_bytes = em_gerar_imagem_google(prompt_final)

                        if imagem_bytes:
                            st.success("Sua imagem foi criada!")
                            st.image(imagem_bytes, caption="Uma reflexão visual do seu sentimento.", use_column_width=True)
                            st.download_button("📥 Baixar Imagem (PNG)", data=imagem_bytes, file_name="espelho_da_mente.png", mime="image/png", use_container_width=True)
                    else:
                        st.error("Não foi possível gerar o prompt para a imagem.")
                except Exception as e:
                    st.error(f"Erro ao gerar a imagem: {e}")

    def page_buscador_vagas():
            st.header("🔎 Buscador de Vagas de Emprego")
            st.markdown("Encontre vagas de emprego usando a Busca Customizada do Google.")

            GSEARCH_KEY = get_api_key('gsearch_key')
            GSEARCH_CX = get_api_key('gsearch_cx')
            if not GSEARCH_KEY or not GSEARCH_CX:
                st.warning("Configure suas chaves de API na aba 'Perfil e Configurações'.")
                return

            def bv_Google_Search(query, api_key, cx):
                try:
                    service = build("customsearch", "v1", developerKey=api_key)
                    res = service.cse().list(q=query, cx=cx, num=10).execute()
                    return res.get("items", [])
                except Exception as e:
                    st.error(f"Erro ao realizar a busca: {e}")
                    return []

            query = st.text_input("Digite a profissão ou área que deseja buscar", placeholder="Ex: Engenheiro de Software Python")
            locais = st.text_input("Deseja restringir por local ou site específico? (Ex: site:linkedin.com OR site:gupy.io)", placeholder="site:linkedin.com OR site:br.indeed.com")
            dork_extra = st.text_input("Termos avançados (opcional, Ex: intitle:vaga OR intext:home office)", placeholder="inurl:emprego OR intitle:oportunidade")

            if st.button("Buscar Vagas", use_container_width=True):
                if query:
                    with st.spinner(f"Buscando vagas para '{query}'..."):
                        full_query_parts = [f'"{query}"'] # Exact match for the profession/area
                        if locais.strip():
                            full_query_parts.append(f"({locais})")
                        else: # Default job boards if no specific sites are provided
                            full_query_parts.append("site:linkedin.com OR site:br.indeed.com OR site:catho.com.br OR site:infojobs.com.br OR site:vagas.com.br OR site:portal.gupy.io OR site:trabalhabrasil.com.br OR site:glassdoor.com.br OR site:empregos.com.br")

                        if dork_extra.strip():
                            full_query_parts.append(dork_extra)

                        final_dork = " ".join(full_query_parts)

                        st.caption(f"Dork gerado: `{final_dork}`")
                        resultados = bv_Google_Search(final_dork, GSEARCH_KEY, GSEARCH_CX)

                        if resultados:
                            st.success(f"🔗 {len(resultados)} resultado(s) encontrados!")
                            for r in resultados:
                                st.markdown(f"### [{r['title']}]({r['link']})")
                                st.caption(f"🌐 Fonte: {r['displayLink']}")
                                st.write(r.get('snippet', 'Sem descrição disponível.'))
                                st.markdown("---")
                        else:
                            st.warning("Nenhum resultado encontrado. Experimente alterar o termo ou o filtro.")
                else:
                    st.warning("Por favor, digite um termo para a busca.")


    def page_meus_arquivos():
            st.header("🗂️ Meus Arquivos")
            st.markdown("Aqui estão os arquivos que você gerou e salvou na plataforma.")

            user_dir = Path(__file__).parent / "user_files" / username
            if not user_dir.exists() or not any(user_dir.iterdir()):
                st.info("Você ainda não salvou nenhum arquivo.")
                return

            files = os.listdir(user_dir)
            if not files: # Check again in case iterdir() returned something but listdir is empty (e.g., hidden files)
                 st.info("Você ainda não salvou nenhum arquivo.")
                 return

            for file_name in files:
                file_path = user_dir / file_name
                with open(file_path, "rb") as f:
                    st.download_button(
                        label=f"📥 Baixar {file_name}",
                        data=f,
                        file_name=file_name,
                        key=f"download_{file_name}"
                    )
                st.markdown("---")

    def page_perfil_configuracoes():
            st.header("👤 Perfil e Configurações")
            st.subheader("Gerencie suas chaves de API")
            st.info("Suas chaves são armazenadas de forma segura e usadas para alimentar as ferramentas da plataforma.")

            with st.form("api_keys_form"):
                st.write("**Chaves do Google / Gemini**")
                gemini_key = st.text_input("Gemini API Key", value=get_api_key('gemini_key') or "", type="password")
                gcp_project_id = st.text_input("Google Cloud Project ID", value=get_api_key('gcp_project_id') or "")
                gcp_location = st.text_input("Google Cloud Location", value=get_api_key('gcp_location') or "")

                st.write("**Chaves do Google Custom Search**")
                gsearch_key = st.text_input("Google Custom Search API Key", value=get_api_key('gsearch_key') or "", type="password")
                gsearch_cx = st.text_input("Google Custom Search CX ID", value=get_api_key('gsearch_cx') or "")

                submitted = st.form_submit_button("Salvar Configurações")
                if submitted:
                    user_credentials = config['credentials']['usernames'][username]
                    user_credentials['api_keys']['gemini_key'] = gemini_key
                    user_credentials['api_keys']['gcp_project_id'] = gcp_project_id
                    user_credentials['api_keys']['gcp_location'] = gcp_location
                    user_credentials['api_keys']['gsearch_key'] = gsearch_key
                    user_credentials['api_keys']['gsearch_cx'] = gsearch_cx

                    with open(config_file, 'w') as file:
                        yaml.dump(config, file, default_flow_style=False)
                    st.success("Suas configurações foram salvas com sucesso!")
                    st.rerun()

    # --- ROTEADOR DE PÁGINAS ---
    if page == "Página Inicial":
        page_inicial()
    elif page == "Gerador de Exercícios":
        page_gerador_exercicios()
    elif page == "Otimizador de Prompt":
        page_otimizador_prompt()
    elif page == "Análise Visual de Imagens":
        page_analise_visual()
    elif page == "Criador de Aplicativos":
        page_criador_aplicativos()
    elif page == "Fábrica de Spritesheets 2D":
        page_fabrica_spritesheets()
    elif page == "Análise de Logs":
        page_analise_logs()
    elif page == "Espelho da Mente":
        page_espelho_da_mente()
    elif page == "Buscador de Vagas":
        page_buscador_vagas()
    elif page == "Meus Arquivos":
        page_meus_arquivos()
    elif page == "Perfil e Configurações":
        page_perfil_configuracoes()
